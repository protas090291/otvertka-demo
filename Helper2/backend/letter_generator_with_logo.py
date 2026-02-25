#!/usr/bin/env python3
"""
Генератор писем с логотипом в точном стиле пользователя
"""

import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class LetterGeneratorWithLogo:
    def __init__(self, documents_dir: str = "../existing_documents"):
        self.documents_dir = documents_dir
        self.logo_path = "logo_image1.png"  # Путь к извлеченному логотипу
        os.makedirs(documents_dir, exist_ok=True)
    
    def create_letter_with_logo(self, apartment_id: str, issue_type: str, issue_description: str) -> str:
        """Создает письмо с логотипом точно в стиле оригинальных писем"""
        
        # Создаем документ
        doc = Document()
        
        # Устанавливаем поля документа как в оригинале
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)
        
        # 1. ТАБЛИЦА С ЛОГОТИПОМ И АДРЕСОМ
        self._add_logo_and_address_table(doc)
        
        # 2. ОБРАЩЕНИЕ
        self._add_greeting(doc)
        
        # 3. ОСНОВНОЙ ТЕКСТ
        self._add_main_content(doc, apartment_id, issue_type, issue_description)
        
        # 4. ПРОСЬБА О СОДЕЙСТВИИ
        self._add_request_for_assistance(doc)
        
        # 5. ПОДПИСЬ
        self._add_signature(doc)
        
        # Сохраняем документ
        filename = f"Исх. письмо {issue_type} {apartment_id} от {datetime.now().strftime('%d.%m.%y')}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_logo_and_address_table(self, doc: Document):
        """Добавляет таблицу с логотипом и адресом"""
        # Создаем таблицу 1x2 как в оригинале
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Получаем ячейки
        cells = table.rows[0].cells
        
        # Левая ячейка - логотип и адрес отправителя
        left_cell = cells[0]
        
        # Добавляем логотип, если он существует
        if os.path.exists(self.logo_path):
            try:
                # Добавляем изображение в ячейку
                paragraph = left_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Добавляем адрес под логотипом
                address_para = left_cell.add_paragraph()
                address_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
                address_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            except Exception as e:
                print(f"Ошибка добавления логотипа: {e}")
                # Если не удалось добавить логотип, добавляем только адрес
                left_para = left_cell.paragraphs[0]
                left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
                left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Если логотип не найден, добавляем только адрес
            left_para = left_cell.paragraphs[0]
            left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Правая ячейка - адрес получателя
        right_cell = cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.text = "Руководителю проекта\nООО «АВ Development»\nЭльман И.И."
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Добавляем пустую строку после таблицы
        doc.add_paragraph()
    
    def _add_greeting(self, doc: Document):
        """Добавляет обращение"""
        greeting = doc.add_paragraph()
        greeting.text = "Уважаемый Иса Исаевич!"
        greeting.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем пустую строку
        doc.add_paragraph()
    
    def _add_main_content(self, doc: Document, apartment_id: str, issue_type: str, issue_description: str):
        """Добавляет основной текст письма"""
        
        # Начало письма
        start_para = doc.add_paragraph()
        start_para.text = f"Уведомляем Вас о том, что {issue_description.lower()}"
        start_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Дополнительная информация в зависимости от типа проблемы
        if "смещение сроков" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_para.text = "Данный факт влияет на сроки производства работ и монтаж инженерных систем, в том числе системы отопления, водоснабжения, вентиляции и электроснабжения компанией ООО «Интербилдинг» - сроки будут увеличены."
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif "дефект" in issue_type.lower() or "проблема" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_para.text = f"Обнаруженные дефекты в квартире {apartment_id} требуют немедленного устранения для обеспечения качества выполняемых работ и соблюдения сроков сдачи объекта."
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Добавляем пустую строку
        doc.add_paragraph()
    
    def _add_request_for_assistance(self, doc: Document):
        """Добавляет просьбу о содействии"""
        request_para = doc.add_paragraph()
        request_para.text = "Просим Вас посодействовать в решении данного вопроса для ускорения процесса сдачи и передачи инженерных систем компанией ООО «Сварго» и выполнения монтажных работ инженерных систем компанией ООО «Интербилдинг»."
        request_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Добавляем несколько пустых строк перед подписью
        for _ in range(8):
            doc.add_paragraph()
    
    def _add_signature(self, doc: Document):
        """Добавляет подпись в стиле оригинала"""
        # Должность
        position_para1 = doc.add_paragraph()
        position_para1.text = "            Заместитель директора и"
        position_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        position_para2 = doc.add_paragraph()
        position_para2.text = "            руководитель проекта строительства"
        position_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Компания и подпись
        signature_para = doc.add_paragraph()
        signature_para.text = "            ООО «Интербилдинг»                                                                                         Кучун Р.В."
        signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def create_letter_with_logo():
    """Создает письмо с логотипом"""
    print("📧 Создание письма с логотипом...")
    
    generator = LetterGeneratorWithLogo()
    
    # Тестовые данные
    apartment_id = "1503"
    issue_type = "техническая проблема"
    issue_description = "обнаружены трещины в стенах после завершения штукатурных работ в квартире 1503, требующие дополнительного обследования"
    
    try:
        result = generator.create_letter_with_logo(
            apartment_id=apartment_id,
            issue_type=issue_type,
            issue_description=issue_description
        )
        
        if result:
            print("✅ Письмо с логотипом создано!")
            print(f"📁 Файл: {result}")
            
            # Показываем содержимое
            print("\n📄 СОДЕРЖИМОЕ ПИСЬМА С ЛОГОТИПОМ:")
            print("=" * 60)
            
            from docx import Document
            doc = Document(result)
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    print(f"{i:2d}. {paragraph.text}")
            
            print("=" * 60)
            
            return result
        else:
            print("❌ Не удалось создать письмо")
            return None
            
    except Exception as e:
        print(f"❌ Ошибка при создании письма: {e}")
        return None

if __name__ == "__main__":
    create_letter_with_logo()



