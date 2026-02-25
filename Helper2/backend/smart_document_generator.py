import os
import uuid
import json
import requests
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class SmartDocumentGenerator:
    def __init__(self, documents_dir: str = "documents", supabase_url: str = None, supabase_key: str = None):
        self.documents_dir = documents_dir
        self.supabase_url = supabase_url
        self.supabase_key = supabase_key
        os.makedirs(documents_dir, exist_ok=True)
    
    def get_supabase_data(self, table: str, filters: Dict[str, Any] = None) -> List[Dict]:
        """Получение данных из Supabase"""
        if not self.supabase_url or not self.supabase_key:
            return []
        
        try:
            url = f"{self.supabase_url}/rest/v1/{table}"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json'
            }
            
            params = {}
            if filters:
                for key, value in filters.items():
                    params[key] = f'eq.{value}'
            
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Ошибка получения данных из Supabase: {e}")
            return []
    
    def get_apartment_defects(self, apartment_id: str) -> List[Dict]:
        """Получение дефектов для квартиры"""
        return self.get_supabase_data('defects', {'apartment_id': apartment_id})
    
    def get_apartment_progress(self, apartment_id: str) -> List[Dict]:
        """Получение прогресса для квартиры"""
        return self.get_supabase_data('progress_data', {'apartment_id': apartment_id})
    
    def get_apartment_work_journal(self, apartment_id: str) -> List[Dict]:
        """Получение журнала работ для квартиры"""
        return self.get_supabase_data('work_journal', {'apartment_id': apartment_id})
    
    def get_recent_work_journal(self, days: int = 7) -> List[Dict]:
        """Получение недавних записей журнала работ"""
        try:
            url = f"{self.supabase_url}/rest/v1/work_journal"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json'
            }
            
            # Получаем записи за последние N дней
            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
            
            params = {
                'work_date': f'gte.{start_date}',
                'work_date': f'lte.{end_date}',
                'order': 'work_date.desc'
            }
            
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Ошибка получения журнала работ: {e}")
            return []
    
    def generate_smart_handover_act(self, command_data: Dict[str, Any]) -> str:
        """Генерирует умный акт приёмки на основе реальных данных"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем реальные данные
        defects = self.get_apartment_defects(apartment_id)
        progress_data = self.get_apartment_progress(apartment_id)
        work_journal = self.get_apartment_work_journal(apartment_id)
        
        # Основная информация
        doc.add_heading('Информация о квартире', level=2)
        
        # Таблица с данными
        table = doc.add_table(rows=8, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Анализируем данные
        total_defects = len(defects)
        active_defects = len([d for d in defects if d.get('status') == 'active'])
        fixed_defects = len([d for d in defects if d.get('status') == 'fixed'])
        
        # Вычисляем общий прогресс
        total_progress = 0
        if progress_data:
            total_progress = sum(item.get('fact_progress', 0) for item in progress_data) / len(progress_data)
        
        # Последние работы
        recent_works = work_journal[:5] if work_journal else []
        
        data = [
            ('Номер квартиры:', apartment_id),
            ('Тип акта:', command_data.get('act_type', 'Приёмка')),
            ('Общий прогресс:', f'{total_progress:.1f}%'),
            ('Всего дефектов:', str(total_defects)),
            ('Активные дефекты:', str(active_defects)),
            ('Исправленные дефекты:', str(fixed_defects)),
            ('Статус приёмки:', 'Готов к приёмке' if active_defects == 0 else 'Требует доработки'),
            ('Примечания:', command_data.get('meta', {}).get('notes', 'Без замечаний'))
        ]
        
        for i, (key, value) in enumerate(data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            # Жирный шрифт для ключей
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Список дефектов
        if defects:
            doc.add_heading('Обнаруженные дефекты', level=2)
            defects_table = doc.add_table(rows=1, cols=4)
            defects_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = defects_table.rows[0].cells
            hdr_cells[0].text = 'Описание'
            hdr_cells[1].text = 'Статус'
            hdr_cells[2].text = 'Дата обнаружения'
            hdr_cells[3].text = 'Координаты'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные дефектов
            for defect in defects:
                row_cells = defects_table.add_row().cells
                row_cells[0].text = defect.get('title', 'Не указано')
                row_cells[1].text = 'Активен' if defect.get('status') == 'active' else 'Исправлен'
                row_cells[2].text = defect.get('created_at', '')[:10] if defect.get('created_at') else 'Не указано'
                row_cells[3].text = f"X:{defect.get('x_coord', 0)}, Y:{defect.get('y_coord', 0)}"
        
        # Последние работы
        if recent_works:
            doc.add_heading('Последние выполненные работы', level=2)
            works_table = doc.add_table(rows=1, cols=4)
            works_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = works_table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Задача'
            hdr_cells[2].text = 'Описание'
            hdr_cells[3].text = 'Прогресс'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные работ
            for work in recent_works:
                row_cells = works_table.add_row().cells
                row_cells[0].text = work.get('work_date', 'Не указано')
                row_cells[1].text = work.get('task_name', 'Не указано')
                row_cells[2].text = work.get('work_description', 'Не указано')[:50] + '...' if len(work.get('work_description', '')) > 50 else work.get('work_description', 'Не указано')
                progress_gain = work.get('progress_after', 0) - work.get('progress_before', 0)
                row_cells[3].text = f"+{progress_gain}%"
        
        # Рекомендации на основе данных
        doc.add_heading('Рекомендации', level=2)
        if active_defects > 0:
            doc.add_paragraph(f'• Устранить {active_defects} активных дефектов перед приёмкой')
        if total_progress < 100:
            doc.add_paragraph(f'• Завершить работы (текущий прогресс: {total_progress:.1f}%)')
        if not recent_works:
            doc.add_paragraph('• Проверить актуальность данных о выполненных работах')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Заказчик: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        doc.add_paragraph('Технадзор: _________________')
        
        # Сохраняем документ
        filename = f"smart_act_handover_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_smart_defect_report(self, command_data: Dict[str, Any]) -> str:
        """Генерирует умный отчет о дефектах на основе реальных данных"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем реальные данные
        defects = self.get_apartment_defects(apartment_id)
        progress_data = self.get_apartment_progress(apartment_id)
        
        # Статистика дефектов
        doc.add_heading('Статистика дефектов', level=2)
        
        total_defects = len(defects)
        active_defects = len([d for d in defects if d.get('status') == 'active'])
        fixed_defects = len([d for d in defects if d.get('status') == 'fixed'])
        
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Всего дефектов:', str(total_defects)),
            ('Активных дефектов:', str(active_defects)),
            ('Исправленных дефектов:', str(fixed_defects)),
            ('Процент исправления:', f'{(fixed_defects/total_defects*100):.1f}%' if total_defects > 0 else '0%')
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            for paragraph in stats_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Детальный список дефектов
        if defects:
            doc.add_heading('Детальный список дефектов', level=2)
            defects_table = doc.add_table(rows=1, cols=5)
            defects_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = defects_table.rows[0].cells
            hdr_cells[0].text = 'Название'
            hdr_cells[1].text = 'Описание'
            hdr_cells[2].text = 'Статус'
            hdr_cells[3].text = 'Дата обнаружения'
            hdr_cells[4].text = 'Координаты'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные дефектов
            for defect in defects:
                row_cells = defects_table.add_row().cells
                row_cells[0].text = defect.get('title', 'Не указано')
                row_cells[1].text = defect.get('description', 'Описание отсутствует')
                row_cells[2].text = 'Активен' if defect.get('status') == 'active' else 'Исправлен'
                row_cells[3].text = defect.get('created_at', '')[:10] if defect.get('created_at') else 'Не указано'
                row_cells[4].text = f"X:{defect.get('x_coord', 0)}, Y:{defect.get('y_coord', 0)}"
        
        # Анализ прогресса
        if progress_data:
            doc.add_heading('Анализ прогресса работ', level=2)
            progress_table = doc.add_table(rows=1, cols=3)
            progress_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = progress_table.rows[0].cells
            hdr_cells[0].text = 'Задача'
            hdr_cells[1].text = 'Секция'
            hdr_cells[2].text = 'Прогресс'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные прогресса
            for progress in progress_data:
                row_cells = progress_table.add_row().cells
                row_cells[0].text = progress.get('task_name', 'Не указано')
                row_cells[1].text = progress.get('section', 'Не указано')
                row_cells[2].text = f"{progress.get('fact_progress', 0)}%"
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        if active_defects > 0:
            doc.add_paragraph(f'• Приоритетно устранить {active_defects} активных дефектов')
        if total_defects > 0 and fixed_defects/total_defects < 0.8:
            doc.add_paragraph('• Ускорить процесс устранения дефектов')
        doc.add_paragraph('• Провести повторную проверку после устранения дефектов')
        doc.add_paragraph('• Обновить статус дефектов в системе')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Обнаружил: _________________')
        doc.add_paragraph('Ответственный: _________________')
        doc.add_paragraph('Технадзор: _________________')
        
        # Сохраняем документ
        filename = f"smart_defect_report_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_smart_work_report(self, command_data: Dict[str, Any]) -> str:
        """Генерирует умный отчет о работах на основе реальных данных"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем реальные данные
        work_journal = self.get_apartment_work_journal(apartment_id)
        progress_data = self.get_apartment_progress(apartment_id)
        recent_works = self.get_recent_work_journal(7)  # За последнюю неделю
        
        # Статистика работ
        doc.add_heading('Статистика работ', level=2)
        
        total_works = len(work_journal)
        total_progress_gained = sum(work.get('progress_after', 0) - work.get('progress_before', 0) for work in work_journal)
        unique_workers = len(set(work.get('worker_name', '') for work in work_journal if work.get('worker_name')))
        unique_tasks = len(set(work.get('task_name', '') for work in work_journal if work.get('task_name')))
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Всего записей о работах:', str(total_works)),
            ('Общий прирост прогресса:', f'{total_progress_gained}%'),
            ('Уникальных исполнителей:', str(unique_workers)),
            ('Уникальных задач:', str(unique_tasks)),
            ('Период отчета:', 'За все время')
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            for paragraph in stats_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Последние работы
        if work_journal:
            doc.add_heading('Последние выполненные работы', level=2)
            works_table = doc.add_table(rows=1, cols=6)
            works_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = works_table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Время'
            hdr_cells[2].text = 'Задача'
            hdr_cells[3].text = 'Описание'
            hdr_cells[4].text = 'Исполнитель'
            hdr_cells[5].text = 'Прогресс'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные работ (показываем последние 10)
            for work in work_journal[:10]:
                row_cells = works_table.add_row().cells
                row_cells[0].text = work.get('work_date', 'Не указано')
                row_cells[1].text = work.get('work_time', 'Не указано')
                row_cells[2].text = work.get('task_name', 'Не указано')
                description = work.get('work_description', 'Не указано')
                row_cells[3].text = description[:30] + '...' if len(description) > 30 else description
                row_cells[4].text = work.get('worker_name', 'Не указано')
                progress_gain = work.get('progress_after', 0) - work.get('progress_before', 0)
                row_cells[5].text = f"{work.get('progress_before', 0)}% → {work.get('progress_after', 0)}% (+{progress_gain}%)"
        
        # Анализ прогресса по задачам
        if progress_data:
            doc.add_heading('Текущий прогресс по задачам', level=2)
            progress_table = doc.add_table(rows=1, cols=4)
            progress_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = progress_table.rows[0].cells
            hdr_cells[0].text = 'Задача'
            hdr_cells[1].text = 'Секция'
            hdr_cells[2].text = 'Фактический прогресс'
            hdr_cells[3].text = 'Плановый прогресс'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные прогресса
            for progress in progress_data:
                row_cells = progress_table.add_row().cells
                row_cells[0].text = progress.get('task_name', 'Не указано')
                row_cells[1].text = progress.get('section', 'Не указано')
                row_cells[2].text = f"{progress.get('fact_progress', 0)}%"
                row_cells[3].text = f"{progress.get('plan_progress', 0)}%"
        
        # Анализ эффективности
        doc.add_heading('Анализ эффективности', level=2)
        if total_works > 0:
            avg_progress_per_work = total_progress_gained / total_works
            doc.add_paragraph(f'• Средний прирост прогресса за работу: {avg_progress_per_work:.1f}%')
        
        if unique_workers > 0:
            works_per_worker = total_works / unique_workers
            doc.add_paragraph(f'• Среднее количество работ на исполнителя: {works_per_worker:.1f}')
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        if total_progress_gained < 50:
            doc.add_paragraph('• Увеличить интенсивность работ')
        if unique_workers < 3:
            doc.add_paragraph('• Рассмотреть привлечение дополнительных исполнителей')
        doc.add_paragraph('• Продолжить мониторинг прогресса')
        doc.add_paragraph('• Регулярно обновлять данные в системе')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Прораб: _________________')
        doc.add_paragraph('Технадзор: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        
        # Сохраняем документ
        filename = f"smart_work_report_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath

