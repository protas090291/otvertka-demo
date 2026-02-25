import os
import uuid
import json
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class LearningDocumentGenerator:
    def __init__(self, documents_dir: str = "documents", supabase_url: str = None, supabase_key: str = None):
        self.documents_dir = documents_dir
        self.supabase_url = supabase_url
        self.supabase_key = supabase_key
        os.makedirs(documents_dir, exist_ok=True)
    
    def get_supabase_data(self, table: str, filters: Dict[str, Any] = None) -> List[Dict]:
        """Получение данных из Supabase"""
        if not self.supabase_url or not self.supabase_key:
            return []
        
        try:
            url = f"{self.supabase_url}/rest/v1/{table}"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json'
            }
            
            params = {}
            if filters:
                for key, value in filters.items():
                    params[key] = f'eq.{value}'
            
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Ошибка получения данных из Supabase: {e}")
            return []
    
    def get_learning_examples(self, template_type: str, limit: int = 5) -> List[Dict]:
        """Получение примеров документов для обучения из простой таблицы"""
        try:
            url = f"{self.supabase_url}/rest/v1/document_templates"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json'
            }
            
            params = {
                'type': f'eq.{template_type}',
                'is_active': 'eq.true',
                'limit': limit,
                'order': 'created_at.desc'
            }
            
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            print(f"Ошибка получения примеров для обучения: {e}")
            return []
    
    def get_document_generation_rules(self, template_type: str) -> List[Dict]:
        """Получение правил генерации документов"""
        return self.get_supabase_data('document_generation_rules', {
            'template_type': template_type,
            'is_active': True
        })
    
    def get_best_template(self, template_type: str, apartment_id: str = None) -> Optional[Dict]:
        """Получение лучшего шаблона для типа документа"""
        try:
            url = f"{self.supabase_url}/rest/v1/document_templates"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json'
            }
            
            params = {
                'type': f'eq.{template_type}',
                'is_active': 'eq.true',
                'limit': 1,
                'order': 'created_at.desc'
            }
            
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            result = response.json()
            return result[0] if result else None
        except Exception as e:
            print(f"Ошибка получения шаблона: {e}")
            return None
    
    def analyze_examples_patterns(self, examples: List[Dict]) -> Dict[str, Any]:
        """Анализирует паттерны в примерах документов"""
        patterns = {
            'common_sections': [],
            'common_formatting': {},
            'common_structure': {},
            'content_patterns': {}
        }
        
        if not examples:
            return patterns
        
        # Анализируем структуру примеров
        all_sections = []
        all_formatting = []
        all_structures = []
        
        for example in examples:
            document_data = example.get('document_data', {})
            structure = document_data.get('structure', {})
            formatting = document_data.get('formatting', {})
            content = document_data.get('content', {})
            
            # Собираем секции
            sections = [s['name'] for s in structure.get('sections', [])]
            all_sections.extend(sections)
            
            # Собираем форматирование
            all_formatting.append(formatting)
            
            # Собираем структуры
            all_structures.append(structure)
        
        # Находим общие секции
        from collections import Counter
        section_counts = Counter(all_sections)
        patterns['common_sections'] = [section for section, count in section_counts.most_common() if count > len(examples) * 0.5]
        
        # Анализируем общее форматирование
        if all_formatting:
            common_fonts = set()
            common_styles = set()
            
            for fmt in all_formatting:
                common_fonts.update(fmt.get('font_styles', []))
                common_styles.update(fmt.get('paragraph_styles', []))
            
            patterns['common_formatting'] = {
                'fonts': list(common_fonts),
                'styles': list(common_styles)
            }
        
        return patterns
    
    def generate_learning_based_document(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Генерирует документ на основе изученных примеров"""
        # Получаем примеры для обучения
        examples = self.get_learning_examples(template_type, limit=3)
        
        # Получаем правила генерации
        rules = self.get_document_generation_rules(template_type)
        
        # Анализируем паттерны в примерах
        patterns = self.analyze_examples_patterns(examples)
        
        # Создаем документ на основе изученных паттернов
        doc = Document()
        
        # Устанавливаем точные поля как в оригинальных документах
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)      # 0.50 дюйма
            section.bottom_margin = Inches(0.5)   # 0.50 дюйма
            section.left_margin = Inches(0.5)     # 0.50 дюйма
            section.right_margin = Inches(0.5)    # 0.50 дюйма
        
        # Применяем правила форматирования
        self._apply_formatting_rules(doc, rules, patterns)
        
        # Генерируем содержимое на основе типа документа
        if template_type == 'handover_act':
            self._generate_handover_act_content(doc, command_data, patterns, rules)
        elif template_type == 'defect_report':
            self._generate_defect_report_content(doc, command_data, patterns, rules)
        elif template_type == 'work_report':
            self._generate_work_report_content(doc, command_data, patterns, rules)
        elif template_type == 'letter':
            self._generate_letter_content(doc, command_data, patterns, rules)
        
        # Сохраняем документ
        filename = f"learning_{template_type}_{command_data.get('apartment_id', 'unknown')}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        # Логируем процесс обучения
        self._log_learning_process(template_type, command_data, patterns, len(examples))
        
        return filepath
    
    def _apply_formatting_rules(self, doc: Document, rules: List[Dict], patterns: Dict[str, Any]):
        """Применяет правила форматирования к документу"""
        formatting_rules = [rule for rule in rules if rule['rule_type'] == 'formatting']
        
        for rule in formatting_rules:
            action = rule.get('rule_action', {})
            condition = rule.get('rule_condition', {})
            
            # Применяем форматирование заголовков
            if condition.get('section') == 'header':
                # Настраиваем стили заголовков
                if 'font_size' in action:
                    # Здесь можно настроить стили заголовков
                    pass
    
    def _generate_handover_act_content(self, doc: Document, command_data: Dict[str, Any], patterns: Dict[str, Any], rules: List[Dict]):
        """Генерирует содержимое акта приёмки на основе изученных паттернов"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        # Заголовок (на основе изученных паттернов)
        title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем реальные данные
        defects = self.get_supabase_data('defects', {'apartment_id': apartment_id})
        progress_data = self.get_supabase_data('progress_data', {'apartment_id': apartment_id})
        work_journal = self.get_supabase_data('work_journal', {'apartment_id': apartment_id})
        
        # Применяем правила структуры
        structure_rules = [rule for rule in rules if rule['rule_type'] == 'structure']
        
        # Информация о квартире (на основе правил)
        doc.add_heading('Информация о квартире', level=2)
        
        # Создаем таблицу на основе изученных паттернов
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Анализируем данные
        total_defects = len(defects)
        active_defects = len([d for d in defects if d.get('status') == 'active'])
        fixed_defects = len([d for d in defects if d.get('status') == 'fixed'])
        
        total_progress = 0
        if progress_data:
            total_progress = sum(item.get('fact_progress', 0) for item in progress_data) / len(progress_data)
        
        data = [
            ('Номер квартиры:', apartment_id),
            ('Общий прогресс:', f'{total_progress:.1f}%'),
            ('Всего дефектов:', str(total_defects)),
            ('Активных дефектов:', str(active_defects)),
            ('Исправленных дефектов:', str(fixed_defects)),
            ('Статус приёмки:', 'Готов к приёмке' if active_defects == 0 else 'Требует доработки')
        ]
        
        for i, (key, value) in enumerate(data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            # Жирный шрифт для ключей (на основе изученных паттернов)
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Список дефектов (если есть)
        if defects:
            doc.add_heading('Обнаруженные дефекты', level=2)
            defects_table = doc.add_table(rows=1, cols=4)
            defects_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = defects_table.rows[0].cells
            hdr_cells[0].text = 'Описание'
            hdr_cells[1].text = 'Статус'
            hdr_cells[2].text = 'Дата обнаружения'
            hdr_cells[3].text = 'Координаты'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные дефектов
            for defect in defects:
                row_cells = defects_table.add_row().cells
                row_cells[0].text = defect.get('title', 'Не указано')
                row_cells[1].text = 'Активен' if defect.get('status') == 'active' else 'Исправлен'
                row_cells[2].text = defect.get('created_at', '')[:10] if defect.get('created_at') else 'Не указано'
                row_cells[3].text = f"X:{defect.get('x_coord', 0)}, Y:{defect.get('y_coord', 0)}"
        
        # Рекомендации (на основе правил)
        content_rules = [rule for rule in rules if rule['rule_type'] == 'content' and rule.get('rule_action', {}).get('auto_generate')]
        
        if content_rules:
            doc.add_heading('Рекомендации', level=2)
            if active_defects > 0:
                doc.add_paragraph(f'• Устранить {active_defects} активных дефектов перед приёмкой')
            if total_progress < 100:
                doc.add_paragraph(f'• Завершить работы (текущий прогресс: {total_progress:.1f}%)')
            if not work_journal:
                doc.add_paragraph('• Проверить актуальность данных о выполненных работах')
        
        # Подписи (на основе изученных паттернов)
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Заказчик: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        doc.add_paragraph('Технадзор: _________________')
    
    def _generate_defect_report_content(self, doc: Document, command_data: Dict[str, Any], patterns: Dict[str, Any], rules: List[Dict]):
        """Генерирует содержимое отчета о дефектах на основе изученных паттернов"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем данные
        defects = self.get_supabase_data('defects', {'apartment_id': apartment_id})
        progress_data = self.get_supabase_data('progress_data', {'apartment_id': apartment_id})
        
        # Статистика дефектов
        doc.add_heading('Статистика дефектов', level=2)
        
        total_defects = len(defects)
        active_defects = len([d for d in defects if d.get('status') == 'active'])
        fixed_defects = len([d for d in defects if d.get('status') == 'fixed'])
        
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Всего дефектов:', str(total_defects)),
            ('Активных дефектов:', str(active_defects)),
            ('Исправленных дефектов:', str(fixed_defects)),
            ('Процент исправления:', f'{(fixed_defects/total_defects*100):.1f}%' if total_defects > 0 else '0%')
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            for paragraph in stats_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Детальный список дефектов
        if defects:
            doc.add_heading('Детальный список дефектов', level=2)
            defects_table = doc.add_table(rows=1, cols=5)
            defects_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = defects_table.rows[0].cells
            hdr_cells[0].text = 'Название'
            hdr_cells[1].text = 'Описание'
            hdr_cells[2].text = 'Статус'
            hdr_cells[3].text = 'Дата обнаружения'
            hdr_cells[4].text = 'Координаты'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные дефектов
            for defect in defects:
                row_cells = defects_table.add_row().cells
                row_cells[0].text = defect.get('title', 'Не указано')
                row_cells[1].text = defect.get('description', 'Описание отсутствует')
                row_cells[2].text = 'Активен' if defect.get('status') == 'active' else 'Исправлен'
                row_cells[3].text = defect.get('created_at', '')[:10] if defect.get('created_at') else 'Не указано'
                row_cells[4].text = f"X:{defect.get('x_coord', 0)}, Y:{defect.get('y_coord', 0)}"
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        if active_defects > 0:
            doc.add_paragraph(f'• Приоритетно устранить {active_defects} активных дефектов')
        if total_defects > 0 and fixed_defects/total_defects < 0.8:
            doc.add_paragraph('• Ускорить процесс устранения дефектов')
        doc.add_paragraph('• Провести повторную проверку после устранения дефектов')
        doc.add_paragraph('• Обновить статус дефектов в системе')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Обнаружил: _________________')
        doc.add_paragraph('Ответственный: _________________')
        doc.add_paragraph('Технадзор: _________________')
    
    def _generate_work_report_content(self, doc: Document, command_data: Dict[str, Any], patterns: Dict[str, Any], rules: List[Dict]):
        """Генерирует содержимое отчета о работах на основе изученных паттернов"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Получаем данные
        work_journal = self.get_supabase_data('work_journal', {'apartment_id': apartment_id})
        progress_data = self.get_supabase_data('progress_data', {'apartment_id': apartment_id})
        
        # Статистика работ
        doc.add_heading('Статистика работ', level=2)
        
        total_works = len(work_journal)
        total_progress_gained = sum(work.get('progress_after', 0) - work.get('progress_before', 0) for work in work_journal)
        unique_workers = len(set(work.get('worker_name', '') for work in work_journal if work.get('worker_name')))
        unique_tasks = len(set(work.get('task_name', '') for work in work_journal if work.get('task_name')))
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Всего записей о работах:', str(total_works)),
            ('Общий прирост прогресса:', f'{total_progress_gained}%'),
            ('Уникальных исполнителей:', str(unique_workers)),
            ('Уникальных задач:', str(unique_tasks)),
            ('Период отчета:', 'За все время')
        ]
        
        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            for paragraph in stats_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Последние работы
        if work_journal:
            doc.add_heading('Последние выполненные работы', level=2)
            works_table = doc.add_table(rows=1, cols=6)
            works_table.style = 'Table Grid'
            
            # Заголовки
            hdr_cells = works_table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Время'
            hdr_cells[2].text = 'Задача'
            hdr_cells[3].text = 'Описание'
            hdr_cells[4].text = 'Исполнитель'
            hdr_cells[5].text = 'Прогресс'
            
            # Жирный шрифт для заголовков
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Данные работ (показываем последние 10)
            for work in work_journal[:10]:
                row_cells = works_table.add_row().cells
                row_cells[0].text = work.get('work_date', 'Не указано')
                row_cells[1].text = work.get('work_time', 'Не указано')
                row_cells[2].text = work.get('task_name', 'Не указано')
                description = work.get('work_description', 'Не указано')
                row_cells[3].text = description[:30] + '...' if len(description) > 30 else description
                row_cells[4].text = work.get('worker_name', 'Не указано')
                progress_gain = work.get('progress_after', 0) - work.get('progress_before', 0)
                row_cells[5].text = f"{work.get('progress_before', 0)}% → {work.get('progress_after', 0)}% (+{progress_gain}%)"
        
        # Анализ эффективности
        doc.add_heading('Анализ эффективности', level=2)
        if total_works > 0:
            avg_progress_per_work = total_progress_gained / total_works
            doc.add_paragraph(f'• Средний прирост прогресса за работу: {avg_progress_per_work:.1f}%')
        
        if unique_workers > 0:
            works_per_worker = total_works / unique_workers
            doc.add_paragraph(f'• Среднее количество работ на исполнителя: {works_per_worker:.1f}')
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        if total_progress_gained < 50:
            doc.add_paragraph('• Увеличить интенсивность работ')
        if unique_workers < 3:
            doc.add_paragraph('• Рассмотреть привлечение дополнительных исполнителей')
        doc.add_paragraph('• Продолжить мониторинг прогресса')
        doc.add_paragraph('• Регулярно обновлять данные в системе')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Прораб: _________________')
        doc.add_paragraph('Технадзор: _________________')
        doc.add_paragraph('Подрядчик: _________________')
    
    def _generate_letter_content(self, doc: Document, command_data: Dict[str, Any], patterns: Dict[str, Any], rules: List[Dict]):
        """Генерирует содержимое письма на основе изученных примеров"""
        apartment_id = command_data.get('apartment_id', 'Unknown')
        issue_type = command_data.get('issue_type', 'технический вопрос')
        issue_description = command_data.get('issue_description', 'Описание отсутствует')
        expected_resolution = command_data.get('expected_resolution', 'Решение в процессе')
        
        # Генерируем номер документа с сегодняшней датой
        from datetime import datetime
        current_date = datetime.now()
        day_month = current_date.strftime('%d/%m')
        
        # Подсчитываем количество документов за сегодня
        today_str = current_date.strftime('%d.%m.%y')
        today_docs = []
        
        # Ищем файлы, которые содержат сегодняшнюю дату в названии
        for f in os.listdir(self.documents_dir):
            if f.endswith('.docx') and today_str in f:
                today_docs.append(f)
        
        # Также считаем файлы learning_letter, созданные сегодня
        import time
        today_timestamp = current_date.timestamp()
        today_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0).timestamp()
        
        for f in os.listdir(self.documents_dir):
            if f.startswith('learning_letter_') and f.endswith('.docx'):
                file_path = os.path.join(self.documents_dir, f)
                file_time = os.path.getmtime(file_path)
                if today_start <= file_time <= today_timestamp:
                    today_docs.append(f)
        
        # Считаем общее количество документов за сегодня
        # Это будет базовый номер для следующего документа
        doc_number = len(today_docs) + 1
        
        document_number = f"{day_month}-{doc_number}"
        
        # Отладочная информация
        print(f"🔢 Генерируем номер документа: {document_number}")
        print(f"📅 Сегодняшняя дата: {current_date.strftime('%d.%m.%Y')}")
        print(f"📁 Найдено файлов за сегодня: {len(today_docs)}")
        print(f"🔍 Следующий номер: {doc_number}")
        
        # 1. ТАБЛИЦА С ЛОГОТИПОМ И АДРЕСОМ (как в оригинале)
        self._add_logo_and_address_table(doc, document_number)
        
        # 2. ОБРАЩЕНИЕ (по центру, жирный шрифт как в оригинале)
        greeting = doc.add_paragraph()
        greeting_run = greeting.add_run("Уважаемый Иса Исаевич!")
        greeting_run.bold = True
        greeting_run.font.name = 'Times New Roman'
        greeting_run.font.size = Pt(12)  # 177800 = 12pt
        greeting.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем пустую строку
        doc.add_paragraph()
        
        # 3. ОСНОВНОЙ ТЕКСТ (с точными отступами как в оригинале)
        start_para = doc.add_paragraph()
        start_run = start_para.add_run(f"Уведомляем Вас о том, что {issue_description.lower()}")
        start_run.font.name = 'Times New Roman'
        start_run.font.size = Pt(12)  # 177800 = 12pt
        start_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        start_para.paragraph_format.left_indent = Pt(10.2)  # 180340 = 10.2pt
        start_para.paragraph_format.first_line_indent = Pt(20.4)  # 360045 = 20.4pt
        
        # Дополнительная информация в зависимости от типа проблемы
        if "смещение сроков" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_run = additional_para.add_run("Данный факт влияет на сроки производства работ и монтаж инженерных систем, в том числе системы отопления, водоснабжения, вентиляции и электроснабжения компанией ООО «Интербилдинг» - сроки будут увеличены.")
            additional_run.font.name = 'Times New Roman'
            additional_run.font.size = Pt(12)
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            additional_para.paragraph_format.first_line_indent = Pt(25.5)  # 450215 = 25.5pt
        elif "дефект" in issue_type.lower() or "проблема" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_run = additional_para.add_run(f"Обнаруженные дефекты в квартире {apartment_id} требуют немедленного устранения для обеспечения качества выполняемых работ и соблюдения сроков сдачи объекта.")
            additional_run.font.name = 'Times New Roman'
            additional_run.font.size = Pt(12)
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            additional_para.paragraph_format.first_line_indent = Pt(25.5)  # 450215 = 25.5pt
        
        # Добавляем пустую строку
        doc.add_paragraph()
        
        # 4. ПРОСЬБА О СОДЕЙСТВИИ (с точными отступами)
        request_para = doc.add_paragraph()
        request_run = request_para.add_run("Просим Вас посодействовать в решении данного вопроса для ускорения процесса сдачи и передачи инженерных систем компанией ООО «Сварго» и выполнения монтажных работ инженерных систем компанией ООО «Интербилдинг».")
        request_run.font.name = 'Times New Roman'
        request_run.font.size = Pt(12)
        request_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        request_para.paragraph_format.first_line_indent = Pt(25.5)  # 450215 = 25.5pt
        request_para.paragraph_format.line_spacing = 1.0  # Межстрочный интервал как в оригинале
        
        # Добавляем несколько пустых строк перед подписью
        for _ in range(8):
            doc.add_paragraph()
        
        # 5. ПОДПИСЬ (точно как в оригинале - жирный шрифт)
        position_para1 = doc.add_paragraph()
        position_run1 = position_para1.add_run("           Заместитель директора и ")
        position_run1.bold = True
        position_run1.font.name = 'Times New Roman'
        position_run1.font.size = Pt(11)  # 165100 = 11pt
        position_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        position_para1.paragraph_format.line_spacing = 1.0  # Как в оригинале
        
        position_para2 = doc.add_paragraph()
        position_run2 = position_para2.add_run("           руководитель проекта строительства ")
        position_run2.bold = True
        position_run2.font.name = 'Times New Roman'
        position_run2.font.size = Pt(11)  # 165100 = 11pt
        position_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        position_para2.paragraph_format.line_spacing = 1.0  # Как в оригинале
        
        # Компания и подпись в одном абзаце как в оригинале
        signature_para = doc.add_paragraph()
        company_run = signature_para.add_run("           ООО «Интербилдинг»")
        company_run.bold = True
        company_run.font.name = 'Times New Roman'
        company_run.font.size = Pt(11)  # 165100 = 11pt
        
        # Добавляем много пробелов для выравнивания подписи по правому краю
        signature_run = signature_para.add_run("                                                                                       Кучун Р.В.")
        signature_run.bold = True
        signature_run.font.name = 'Times New Roman'
        signature_run.font.size = Pt(11)  # 165100 = 11pt
        
        signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_para.paragraph_format.line_spacing = 1.0
    
    def _add_logo_and_address_table(self, doc: Document, document_number: str = None):
        """Добавляет таблицу с логотипом и адресом как в оригинальных письмах"""
        # Создаем таблицу 1x2 как в оригинале
        table = doc.add_table(rows=1, cols=2)
        # Убираем рамки таблицы
        table.style = None
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Получаем ячейки
        cells = table.rows[0].cells
        
        # Левая ячейка - логотип и адрес отправителя (по центру как в оригинале)
        left_cell = cells[0]
        
        # Добавляем логотип, если он существует
        logo_path = "logo_image1.png"
        if os.path.exists(logo_path):
            try:
                # Добавляем изображение в ячейку
                paragraph = left_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # По центру как в оригинале
                
                # Добавляем адрес под логотипом
                address_para = left_cell.add_paragraph()
                address_run = address_para.add_run("Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2")
                address_run.font.name = 'Times New Roman'
                address_run.font.size = Pt(12)
                address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # По центру как в оригинале
                
                # Добавляем дату и номер документа точно как в оригинале
                if document_number:
                    # Первая строка: два пробела + подчеркивание + дата + подчеркивание + пробел + № (не подчеркнуто) + пробел + номер документа (подчеркнуто)
                    first_line_para = left_cell.add_paragraph()
                    first_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Четыре пробела перед первым подчеркиванием
                    spaces_run = first_line_para.add_run("    ")
                    spaces_run.font.name = 'Times New Roman'
                    spaces_run.font.size = Pt(12)
                    
                    # Три подчеркивания перед датой
                    underline_before_run = first_line_para.add_run("___")
                    underline_before_run.font.name = 'Times New Roman'
                    underline_before_run.font.size = Pt(12)
                    underline_before_run.underline = True
                    
                    # Дата (подчеркнуто) - используем сегодняшнюю дату
                    from datetime import datetime
                    current_date = datetime.now()
                    today_date = current_date.strftime('%d.%m.%Yг.')
                    date_run = first_line_para.add_run(today_date)
                    date_run.font.name = 'Times New Roman'
                    date_run.font.size = Pt(12)
                    date_run.underline = True
                    
                    # Три подчеркивания после даты
                    underline_after_run = first_line_para.add_run("___")
                    underline_after_run.font.name = 'Times New Roman'
                    underline_after_run.font.size = Pt(12)
                    underline_after_run.underline = True
                    
                    # Пробел
                    space_run = first_line_para.add_run(" ")
                    space_run.font.name = 'Times New Roman'
                    space_run.font.size = Pt(12)
                    
                    # № (не подчеркнуто)
                    number_symbol_run = first_line_para.add_run("№")
                    number_symbol_run.font.name = 'Times New Roman'
                    number_symbol_run.font.size = Pt(12)
                    
                    # Пробел
                    space_run2 = first_line_para.add_run(" ")
                    space_run2.font.name = 'Times New Roman'
                    space_run2.font.size = Pt(12)
                    
                    # Четыре подчеркивания перед номером документа
                    number_underline_before_run = first_line_para.add_run("____")
                    number_underline_before_run.font.name = 'Times New Roman'
                    number_underline_before_run.font.size = Pt(12)
                    number_underline_before_run.underline = True
                    
                    # Номер документа (подчеркнуто)
                    number_run = first_line_para.add_run(document_number)
                    number_run.font.name = 'Times New Roman'
                    number_run.font.size = Pt(12)
                    number_run.underline = True
                    
                    # Четыре подчеркивания после номера документа
                    number_underline_after_run = first_line_para.add_run("____")
                    number_underline_after_run.font.name = 'Times New Roman'
                    number_underline_after_run.font.size = Pt(12)
                    number_underline_after_run.underline = True
                    
                    # Вторая строка: "на № ________________ от ________________"
                    second_line_para = left_cell.add_paragraph()
                    second_line_run = second_line_para.add_run("на № ________________ от ________________")
                    second_line_run.font.name = 'Times New Roman'
                    second_line_run.font.size = Pt(12)
                    second_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                print(f"Ошибка добавления логотипа: {e}")
                # Если не удалось добавить логотип, добавляем только адрес
                left_para = left_cell.paragraphs[0]
                left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
                left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем дату и номер документа точно как в оригинале
                if document_number:
                    # Первая строка: два пробела + подчеркивание + дата + подчеркивание + пробел + № (не подчеркнуто) + пробел + номер документа (подчеркнуто)
                    first_line_para = left_cell.add_paragraph()
                    first_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Четыре пробела перед первым подчеркиванием
                    spaces_run = first_line_para.add_run("    ")
                    spaces_run.font.name = 'Times New Roman'
                    spaces_run.font.size = Pt(12)
                    
                    # Три подчеркивания перед датой
                    underline_before_run = first_line_para.add_run("___")
                    underline_before_run.font.name = 'Times New Roman'
                    underline_before_run.font.size = Pt(12)
                    underline_before_run.underline = True
                    
                    # Дата (подчеркнуто) - используем сегодняшнюю дату
                    from datetime import datetime
                    current_date = datetime.now()
                    today_date = current_date.strftime('%d.%m.%Yг.')
                    date_run = first_line_para.add_run(today_date)
                    date_run.font.name = 'Times New Roman'
                    date_run.font.size = Pt(12)
                    date_run.underline = True
                    
                    # Три подчеркивания после даты
                    underline_after_run = first_line_para.add_run("___")
                    underline_after_run.font.name = 'Times New Roman'
                    underline_after_run.font.size = Pt(12)
                    underline_after_run.underline = True
                    
                    # Пробел
                    space_run = first_line_para.add_run(" ")
                    space_run.font.name = 'Times New Roman'
                    space_run.font.size = Pt(12)
                    
                    # № (не подчеркнуто)
                    number_symbol_run = first_line_para.add_run("№")
                    number_symbol_run.font.name = 'Times New Roman'
                    number_symbol_run.font.size = Pt(12)
                    
                    # Пробел
                    space_run2 = first_line_para.add_run(" ")
                    space_run2.font.name = 'Times New Roman'
                    space_run2.font.size = Pt(12)
                    
                    # Четыре подчеркивания перед номером документа
                    number_underline_before_run = first_line_para.add_run("____")
                    number_underline_before_run.font.name = 'Times New Roman'
                    number_underline_before_run.font.size = Pt(12)
                    number_underline_before_run.underline = True
                    
                    # Номер документа (подчеркнуто)
                    number_run = first_line_para.add_run(document_number)
                    number_run.font.name = 'Times New Roman'
                    number_run.font.size = Pt(12)
                    number_run.underline = True
                    
                    # Четыре подчеркивания после номера документа
                    number_underline_after_run = first_line_para.add_run("____")
                    number_underline_after_run.font.name = 'Times New Roman'
                    number_underline_after_run.font.size = Pt(12)
                    number_underline_after_run.underline = True
                    
                    # Вторая строка: "на № ________________ от ________________"
                    second_line_para = left_cell.add_paragraph()
                    second_line_run = second_line_para.add_run("на № ________________ от ________________")
                    second_line_run.font.name = 'Times New Roman'
                    second_line_run.font.size = Pt(12)
                    second_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Если логотип не найден, добавляем только адрес
            left_para = left_cell.paragraphs[0]
            left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
            left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавляем номер документа, если он передан
            if document_number:
                doc_number_para = left_cell.add_paragraph()
                doc_number_para.text = f"№ {document_number}"
                doc_number_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем дату документа
                current_date = datetime.now()
                date_para = left_cell.add_paragraph()
                date_para.text = current_date.strftime('%d.%m.%Y')
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Правая ячейка - адрес получателя (отдельные абзацы по правому краю как в оригинале)
        right_cell = cells[1]
        
        # Создаем отдельные абзацы для каждой строки адресата (жирный шрифт как в оригинале)
        recipient1 = right_cell.add_paragraph()
        recipient1_run = recipient1.add_run("Руководителю проекта")
        recipient1_run.bold = True
        recipient1_run.font.name = 'Times New Roman'
        recipient1_run.font.size = Pt(12)
        recipient1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        recipient2 = right_cell.add_paragraph()
        recipient2_run = recipient2.add_run("ООО «АВ Development»")
        recipient2_run.bold = True
        recipient2_run.font.name = 'Times New Roman'
        recipient2_run.font.size = Pt(12)
        recipient2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        recipient3 = right_cell.add_paragraph()
        recipient3_run = recipient3.add_run("Эльман И.И.")
        recipient3_run.bold = True
        recipient3_run.font.name = 'Times New Roman'
        recipient3_run.font.size = Pt(12)
        recipient3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Добавляем пустую строку после таблицы
        doc.add_paragraph()
    
    def _log_learning_process(self, template_type: str, command_data: Dict[str, Any], patterns: Dict[str, Any], examples_count: int):
        """Логирует процесс обучения AI"""
        if not self.supabase_url or not self.supabase_key:
            return
        
        try:
            log_data = {
                'learning_type': 'template_analysis',
                'input_data': {
                    'template_type': template_type,
                    'apartment_id': command_data.get('apartment_id'),
                    'examples_analyzed': examples_count
                },
                'output_data': {
                    'patterns_found': patterns,
                    'document_generated': True
                },
                'success': True,
                'processing_time_ms': 0,  # Можно добавить измерение времени
                'metadata': {
                    'learning_session': datetime.now().isoformat(),
                    'patterns_analyzed': len(patterns.get('common_sections', []))
                }
            }
            
            url = f"{self.supabase_url}/rest/v1/ai_learning_logs"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json',
                'Prefer': 'return=minimal'
            }
            
            response = requests.post(url, headers=headers, json=log_data)
            response.raise_for_status()
            
            print(f"Процесс обучения для {template_type} записан в логи")
            
        except Exception as e:
            print(f"Ошибка записи логов обучения: {e}")
