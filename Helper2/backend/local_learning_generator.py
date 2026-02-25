#!/usr/bin/env python3
"""
Локальный генератор документов на основе существующих файлов
Работает без Supabase - использует локальные файлы как примеры
"""

import os
import uuid
import random
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class LocalLearningGenerator:
    def __init__(self, documents_dir: str = "documents", examples_dir: str = "../existing_documents"):
        self.documents_dir = documents_dir
        self.examples_dir = examples_dir
        os.makedirs(documents_dir, exist_ok=True)
    
    def get_local_examples(self, template_type: str) -> List[str]:
        """Получает список локальных примеров документов"""
        examples = []
        
        if not os.path.exists(self.examples_dir):
            return examples
        
        for filename in os.listdir(self.examples_dir):
            if not filename.endswith('.docx'):
                continue
            
            # Определяем тип документа по имени файла
            doc_type = "handover_act"
            if "дефект" in filename.lower() or "defect" in filename.lower():
                doc_type = "defect_report"
            elif "работа" in filename.lower() or "work" in filename.lower():
                doc_type = "work_report"
            elif "акт" in filename.lower() or "act" in filename.lower():
                doc_type = "handover_act"
            
            if doc_type == template_type:
                examples.append(os.path.join(self.examples_dir, filename))
        
        return examples
    
    def analyze_example_structure(self, file_path: str) -> Dict[str, Any]:
        """Анализирует структуру примера документа"""
        try:
            doc = Document(file_path)
            
            structure = {
                'headings': [],
                'tables': [],
                'paragraphs': len(doc.paragraphs),
                'styles': set()
            }
            
            # Анализируем заголовки
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    structure['headings'].append({
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name
                    })
                structure['styles'].add(paragraph.style.name)
            
            # Анализируем таблицы
            for table in doc.tables:
                table_info = {
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'headers': []
                }
                
                if table.rows:
                    first_row = table.rows[0]
                    for cell in first_row.cells:
                        table_info['headers'].append(cell.text.strip())
                
                structure['tables'].append(table_info)
            
            structure['styles'] = list(structure['styles'])
            return structure
            
        except Exception as e:
            print(f"Ошибка анализа {file_path}: {e}")
            return {}
    
    def generate_learning_based_document(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Генерирует документ на основе локальных примеров"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        # Получаем примеры
        examples = self.get_local_examples(template_type)
        
        if not examples:
            print(f"❌ Не найдено примеров для типа {template_type}")
            return self._generate_fallback_document(template_type, command_data)
        
        # Анализируем первый пример
        example_structure = self.analyze_example_structure(examples[0])
        
        print(f"📚 Используем пример: {os.path.basename(examples[0])}")
        print(f"📊 Структура: {len(example_structure.get('headings', []))} заголовков, {len(example_structure.get('tables', []))} таблиц")
        
        # Создаем документ на основе примера
        doc = Document()
        
        # Применяем структуру из примера
        self._apply_example_structure(doc, example_structure, template_type, apartment_id)
        
        # Сохраняем документ
        filename = f"local_learning_{template_type}_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        print(f"✅ Создан документ на основе примера: {filename}")
        return filepath
    
    def _apply_example_structure(self, doc: Document, structure: Dict[str, Any], template_type: str, apartment_id: str):
        """Применяет структуру из примера к новому документу"""
        
        # Заголовок документа
        if template_type == 'handover_act':
            title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        elif template_type == 'defect_report':
            title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        elif template_type == 'work_report':
            title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        else:
            title = doc.add_heading('ДОКУМЕНТ', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Применяем структуру из примера
        headings = structure.get('headings', [])
        tables = structure.get('tables', [])
        
        # Добавляем секции на основе примера
        for heading in headings[:3]:  # Берем первые 3 заголовка
            if heading['text'] and len(heading['text']) > 3:
                doc.add_heading(heading['text'], level=2)
                
                # Добавляем содержимое секции
                if 'информация' in heading['text'].lower():
                    self._add_apartment_info(doc, apartment_id)
                elif 'дефект' in heading['text'].lower():
                    self._add_defects_section(doc, apartment_id)
                elif 'работа' in heading['text'].lower():
                    self._add_works_section(doc, apartment_id)
                elif 'рекомендация' in heading['text'].lower():
                    self._add_recommendations_section(doc)
                else:
                    doc.add_paragraph(f'Содержимое секции "{heading["text"]}" будет добавлено здесь.')
        
        # Добавляем таблицы на основе примера
        for i, table_info in enumerate(tables[:2]):  # Берем первые 2 таблицы
            if table_info['headers']:
                doc.add_heading(f'Таблица {i+1}', level=2)
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=len(table_info['headers']))
                table.style = 'Table Grid'
                
                # Заголовки
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(table_info['headers']):
                    hdr_cells[j].text = header
                    # Жирный шрифт для заголовков
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Добавляем примеры данных
                if 'квартир' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = apartment_id
                    row_cells[1].text = 'Готов к приёмке' if i == 0 else '100%'
                elif 'дефект' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Проверка качества'
                    row_cells[1].text = 'Активен'
                elif 'работа' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Покраска стен'
                    row_cells[1].text = 'Завершено'
        
        # Подписи (стандартная секция)
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Заказчик: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        doc.add_paragraph('Технадзор: _________________')
    
    def _add_apartment_info(self, doc: Document, apartment_id: str):
        """Добавляет информацию о квартире"""
        doc.add_paragraph(f'Номер квартиры: {apartment_id}')
        doc.add_paragraph(f'Статус: Готов к приёмке')
        doc.add_paragraph(f'Площадь: 45.2 м²')
        doc.add_paragraph(f'Этаж: 11')
    
    def _add_defects_section(self, doc: Document, apartment_id: str):
        """Добавляет секцию дефектов"""
        doc.add_paragraph('Обнаруженные дефекты:')
        doc.add_paragraph('• Трещина в стене (координаты: X:150, Y:200)')
        doc.add_paragraph('• Неровность пола (координаты: X:300, Y:100)')
        doc.add_paragraph('• Проблемы с проводкой (координаты: X:50, Y:250)')
    
    def _add_works_section(self, doc: Document, apartment_id: str):
        """Добавляет секцию работ"""
        doc.add_paragraph('Выполненные работы:')
        doc.add_paragraph('• Покраска стен - завершено')
        doc.add_paragraph('• Укладка плитки - завершено')
        doc.add_paragraph('• Установка сантехники - в процессе')
    
    def _add_recommendations_section(self, doc: Document):
        """Добавляет секцию рекомендаций"""
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph('• Устранить обнаруженные дефекты')
        doc.add_paragraph('• Провести финальную проверку качества')
        doc.add_paragraph('• Подготовить документы для приёмки')
    
    def _generate_fallback_document(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Создает базовый документ если нет примеров"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        doc = Document()
        
        # Базовый заголовок
        if template_type == 'handover_act':
            title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        elif template_type == 'defect_report':
            title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        elif template_type == 'work_report':
            title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        else:
            title = doc.add_heading('ДОКУМЕНТ', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Базовая информация
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('Документ создан на основе базового шаблона.')
        
        # Сохраняем
        filename = f"fallback_{template_type}_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath

def main():
    """Тестирование локального генератора"""
    generator = LocalLearningGenerator()
    
    # Тестовые данные
    test_data = {
        'apartment_id': '1101',
        'act_type': 'handover'
    }
    
    print("🧪 Тестируем локальный генератор документов")
    
    # Тестируем разные типы
    for doc_type in ['handover_act', 'defect_report', 'work_report']:
        print(f"\n📄 Создаем {doc_type}...")
        result = generator.generate_learning_based_document(doc_type, test_data)
        print(f"✅ Результат: {result}")

if __name__ == "__main__":
    main()

