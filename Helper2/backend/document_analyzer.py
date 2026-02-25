import os
import json
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
import requests
import re

class DocumentAnalyzer:
    def __init__(self, supabase_url: str = None, supabase_key: str = None):
        self.supabase_url = supabase_url
        self.supabase_key = supabase_key
    
    def analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """Анализирует структуру документа и извлекает метаданные"""
        try:
            doc = Document(file_path)
            
            analysis = {
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'analysis_date': datetime.now().isoformat(),
                'structure': self._analyze_document_structure(doc),
                'content': self._extract_content_structure(doc),
                'tables': self._analyze_tables(doc),
                'formatting': self._analyze_formatting(doc),
                'metadata': self._extract_metadata(doc)
            }
            
            return analysis
            
        except Exception as e:
            print(f"Ошибка анализа документа {file_path}: {e}")
            return {}
    
    def _analyze_document_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """Анализирует общую структуру документа"""
        structure = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'headings': [],
            'sections': []
        }
        
        # Анализ заголовков
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                structure['headings'].append({
                    'level': paragraph.style.name,
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name
                })
        
        # Определение секций
        current_section = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Определяем секции по ключевым словам
            if any(keyword in text.lower() for keyword in ['информация', 'данные', 'сведения']):
                current_section = 'information'
            elif any(keyword in text.lower() for keyword in ['дефект', 'недостаток', 'проблема']):
                current_section = 'defects'
            elif any(keyword in text.lower() for keyword in ['работа', 'выполнено', 'прогресс']):
                current_section = 'works'
            elif any(keyword in text.lower() for keyword in ['рекомендация', 'совет', 'предложение']):
                current_section = 'recommendations'
            elif any(keyword in text.lower() for keyword in ['подпись', 'подтверждение']):
                current_section = 'signatures'
            
            if current_section and current_section not in [s['name'] for s in structure['sections']]:
                structure['sections'].append({
                    'name': current_section,
                    'start_paragraph': len(structure['sections']),
                    'content_type': 'text'
                })
        
        return structure
    
    def _extract_content_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """Извлекает структурированное содержимое документа"""
        content = {
            'apartment_info': {},
            'defects': [],
            'works': [],
            'recommendations': [],
            'signatures': [],
            'statistics': {}
        }
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Извлечение информации о квартире
            if 'квартир' in text.lower():
                apartment_match = re.search(r'квартир[аы]?\s*№?\s*(\d+)', text, re.IGNORECASE)
                if apartment_match:
                    content['apartment_info']['apartment_id'] = apartment_match.group(1)
            
            # Извлечение даты
            date_match = re.search(r'(\d{1,2}[./]\d{1,2}[./]\d{2,4})', text)
            if date_match:
                content['apartment_info']['date'] = date_match.group(1)
            
            # Извлечение дефектов
            if any(keyword in text.lower() for keyword in ['дефект', 'недостаток', 'проблема', 'трещина']):
                content['defects'].append({
                    'description': text,
                    'type': 'defect',
                    'paragraph_style': paragraph.style.name
                })
            
            # Извлечение работ
            if any(keyword in text.lower() for keyword in ['работа', 'выполнено', 'завершено', 'покраска', 'укладка']):
                content['works'].append({
                    'description': text,
                    'type': 'work',
                    'paragraph_style': paragraph.style.name
                })
            
            # Извлечение рекомендаций
            if any(keyword in text.lower() for keyword in ['рекомендация', 'совет', 'предложение', 'необходимо']):
                content['recommendations'].append({
                    'description': text,
                    'type': 'recommendation',
                    'paragraph_style': paragraph.style.name
                })
            
            # Извлечение подписей
            if any(keyword in text.lower() for keyword in ['подпись', 'подтверждение', 'согласовано']):
                content['signatures'].append({
                    'description': text,
                    'type': 'signature',
                    'paragraph_style': paragraph.style.name
                })
        
        return content
    
    def _analyze_tables(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """Анализирует таблицы в документе"""
        tables_analysis = []
        
        for i, table in enumerate(doc.tables):
            table_info = {
                'table_index': i,
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
                'headers': [],
                'data_types': [],
                'content_summary': ''
            }
            
            if table.rows:
                # Анализ заголовков
                first_row = table.rows[0]
                for cell in first_row.cells:
                    header_text = cell.text.strip()
                    table_info['headers'].append(header_text)
                
                # Анализ типов данных
                for col_idx in range(len(table_info['headers'])):
                    col_data = []
                    for row in table.rows[1:]:  # Пропускаем заголовок
                        if col_idx < len(row.cells):
                            cell_text = row.cells[col_idx].text.strip()
                            col_data.append(cell_text)
                    
                    # Определяем тип данных в колонке
                    data_type = self._determine_column_type(col_data)
                    table_info['data_types'].append(data_type)
                
                # Краткое описание содержимого
                table_info['content_summary'] = self._summarize_table_content(table)
            
            tables_analysis.append(table_info)
        
        return tables_analysis
    
    def _determine_column_type(self, data: List[str]) -> str:
        """Определяет тип данных в колонке"""
        if not data:
            return 'empty'
        
        # Проверяем на числа
        numeric_count = sum(1 for item in data if re.match(r'^\d+%?$', item))
        if numeric_count > len(data) * 0.7:
            return 'numeric'
        
        # Проверяем на даты
        date_count = sum(1 for item in data if re.match(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', item))
        if date_count > len(data) * 0.7:
            return 'date'
        
        # Проверяем на статусы
        status_keywords = ['активен', 'исправлен', 'выполнено', 'завершено', 'в процессе']
        status_count = sum(1 for item in data if any(keyword in item.lower() for keyword in status_keywords))
        if status_count > len(data) * 0.5:
            return 'status'
        
        return 'text'
    
    def _summarize_table_content(self, table: Table) -> str:
        """Создает краткое описание содержимого таблицы"""
        if not table.rows:
            return 'Пустая таблица'
        
        first_row = table.rows[0]
        headers = [cell.text.strip() for cell in first_row.cells]
        
        # Определяем тип таблицы по заголовкам
        if any('дефект' in header.lower() for header in headers):
            return 'Таблица дефектов'
        elif any('работа' in header.lower() for header in headers):
            return 'Таблица работ'
        elif any('прогресс' in header.lower() for header in headers):
            return 'Таблица прогресса'
        elif any('статистика' in header.lower() for header in headers):
            return 'Таблица статистики'
        else:
            return f'Таблица с колонками: {", ".join(headers[:3])}'
    
    def _analyze_formatting(self, doc: DocumentType) -> Dict[str, Any]:
        """Анализирует форматирование документа"""
        formatting = {
            'font_styles': set(),
            'paragraph_styles': set(),
            'alignment_styles': set(),
            'bold_count': 0,
            'italic_count': 0,
            'underline_count': 0
        }
        
        for paragraph in doc.paragraphs:
            formatting['paragraph_styles'].add(paragraph.style.name)
            
            for run in paragraph.runs:
                if run.bold:
                    formatting['bold_count'] += 1
                if run.italic:
                    formatting['italic_count'] += 1
                if run.underline:
                    formatting['underline_count'] += 1
                
                if run.font.name:
                    formatting['font_styles'].add(run.font.name)
        
        # Преобразуем sets в lists для JSON сериализации
        formatting['font_styles'] = list(formatting['font_styles'])
        formatting['paragraph_styles'] = list(formatting['paragraph_styles'])
        formatting['alignment_styles'] = list(formatting['alignment_styles'])
        
        return formatting
    
    def _extract_metadata(self, doc: DocumentType) -> Dict[str, Any]:
        """Извлекает метаданные документа"""
        metadata = {
            'title': '',
            'author': '',
            'created': '',
            'modified': '',
            'keywords': []
        }
        
        # Попытка извлечь метаданные из свойств документа
        try:
            core_props = doc.core_properties
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            metadata['created'] = str(core_props.created) if core_props.created else ''
            metadata['modified'] = str(core_props.modified) if core_props.modified else ''
            metadata['keywords'] = core_props.keywords.split(',') if core_props.keywords else []
        except:
            pass
        
        return metadata
    
    def determine_document_type(self, analysis: Dict[str, Any]) -> str:
        """Определяет тип документа на основе анализа"""
        content = analysis.get('content', {})
        structure = analysis.get('structure', {})
        
        # Анализируем содержимое для определения типа
        defects_count = len(content.get('defects', []))
        works_count = len(content.get('works', []))
        recommendations_count = len(content.get('recommendations', []))
        
        # Проверяем заголовки
        headings_text = ' '.join([h['text'].lower() for h in structure.get('headings', [])])
        
        if 'акт' in headings_text and 'приёмк' in headings_text:
            return 'handover_act'
        elif 'дефект' in headings_text or defects_count > 2:
            return 'defect_report'
        elif 'работа' in headings_text or 'отчет' in headings_text or works_count > 2:
            return 'work_report'
        elif 'прогресс' in headings_text:
            return 'progress_report'
        else:
            return 'unknown'
    
    def upload_analysis_to_supabase(self, analysis: Dict[str, Any], file_path: str, template_type: str = None) -> bool:
        """Загружает анализ документа в Supabase"""
        if not self.supabase_url or not self.supabase_key:
            print("Supabase не настроен")
            return False
        
        try:
            # Определяем тип документа если не указан
            if not template_type:
                template_type = self.determine_document_type(analysis)
            
            # Создаем запись в document_examples
            example_data = {
                'example_name': analysis['file_name'],
                'example_type': template_type,
                'file_name': analysis['file_name'],
                'file_size': analysis['file_size'],
                'document_data': analysis,
                'quality_score': self._assess_document_quality(analysis),
                'is_approved': True,
                'metadata': {
                    'analysis_date': analysis['analysis_date'],
                    'auto_analyzed': True
                }
            }
            
            # Отправляем в Supabase
            url = f"{self.supabase_url}/rest/v1/document_examples"
            headers = {
                'apikey': self.supabase_key,
                'Authorization': f'Bearer {self.supabase_key}',
                'Content-Type': 'application/json',
                'Prefer': 'return=minimal'
            }
            
            response = requests.post(url, headers=headers, json=example_data)
            response.raise_for_status()
            
            print(f"Анализ документа {analysis['file_name']} загружен в Supabase")
            return True
            
        except Exception as e:
            print(f"Ошибка загрузки анализа в Supabase: {e}")
            return False
    
    def _assess_document_quality(self, analysis: Dict[str, Any]) -> int:
        """Оценивает качество документа (1-5)"""
        score = 3  # Базовая оценка
        
        structure = analysis.get('structure', {})
        content = analysis.get('content', {})
        
        # Бонусы за хорошую структуру
        if len(structure.get('headings', [])) > 2:
            score += 1
        
        if len(structure.get('sections', [])) > 3:
            score += 1
        
        # Бонусы за наличие таблиц
        if analysis.get('tables'):
            score += 1
        
        # Бонусы за структурированное содержимое
        if content.get('apartment_info'):
            score += 1
        
        # Ограничиваем максимальную оценку
        return min(score, 5)
    
    def batch_analyze_documents(self, documents_dir: str) -> List[Dict[str, Any]]:
        """Анализирует все документы в папке"""
        results = []
        
        if not os.path.exists(documents_dir):
            print(f"Папка {documents_dir} не существует")
            return results
        
        for filename in os.listdir(documents_dir):
            if filename.endswith(('.docx', '.doc')):
                file_path = os.path.join(documents_dir, filename)
                print(f"Анализируем {filename}...")
                
                analysis = self.analyze_document_structure(file_path)
                if analysis:
                    results.append(analysis)
                    
                    # Загружаем в Supabase
                    self.upload_analysis_to_supabase(analysis, file_path)
        
        return results

