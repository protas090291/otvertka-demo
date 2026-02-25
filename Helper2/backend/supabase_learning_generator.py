#!/usr/bin/env python3
"""
Генератор документов на основе примеров из Supabase Storage
Использует существующие документы в Documents-base как шаблоны
"""

import os
import uuid
import requests
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class SupabaseLearningGenerator:
    def __init__(self, documents_dir: str = "documents"):
        self.documents_dir = documents_dir
        os.makedirs(documents_dir, exist_ok=True)
        
        # Настройки Supabase из .env
        self.supabase_url = None
        self.supabase_key = None
        
        if os.path.exists('.env'):
            with open('.env', 'r', encoding='utf-8') as f:
                for line in f:
                    if line.startswith('SUPABASE_URL='):
                        self.supabase_url = line.split('=', 1)[1].strip()
                    elif line.startswith('SUPABASE_SERVICE_ROLE_KEY='):
                        self.supabase_key = line.split('=', 1)[1].strip()
    
    def get_examples_from_supabase(self, template_type: str) -> List[str]:
        """Получает список примеров документов из Supabase Storage"""
        examples = []
        
        if not self.supabase_url:
            print("❌ Supabase URL не настроен")
            return examples
        
        # Список реальных документов в Documents-base (названия от 1 до 10)
        known_documents = {
            'handover_act': [
                '7.docx',  # Основной профессиональный шаблон (110KB)
                '4.docx',  # Резервный шаблон
                '5.docx',  # Резервный шаблон
            ],
            'defect_report': [
                '7.docx',  # Используем профессиональный шаблон для всех типов
                '9.docx',  # Резервный шаблон
                '10.docx'  # Резервный шаблон
            ],
            'work_report': [
                '7.docx',  # Используем профессиональный шаблон для всех типов
                '2.pdf',   # PDF отчёт
                '3.pdf'    # PDF отчёт
            ],
            'general': [
                '7.docx',  # Основной профессиональный шаблон
                '4.docx', '5.docx', '6.docx', '8.docx', '9.docx', '10.docx'
            ]
        }
        
        # Возвращаем примеры для указанного типа
        return known_documents.get(template_type, [])
    
    def download_example_from_supabase(self, file_path: str) -> Optional[str]:
        """Скачивает пример документа из Supabase Storage"""
        try:
            # Пробуем разные варианты URL для простых названий файлов
            urls_to_try = [
                f"{self.supabase_url}/storage/v1/object/public/Documents-base/{file_path}",
                f"{self.supabase_url}/storage/v1/object/Documents-base/{file_path}",
                f"{self.supabase_url}/storage/v1/object/public/Documents-base/templates/{file_path}",
                f"{self.supabase_url}/storage/v1/object/Documents-base/templates/{file_path}"
            ]
            
            response = None
            working_url = None
            
            for url in urls_to_try:
                try:
                    response = requests.get(url)
                    if response.status_code == 200:
                        working_url = url
                        break
                except:
                    continue
            
            if not working_url:
                raise Exception(f"Не удалось найти файл по ни одному из URL")
            
            # Сохраняем временно для анализа
            temp_filename = f"temp_{uuid.uuid4().hex[:8]}.docx"
            temp_path = os.path.join(self.documents_dir, temp_filename)
            
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            
            print(f"📥 Скачан пример: {os.path.basename(file_path)}")
            return temp_path
            
        except Exception as e:
            print(f"❌ Ошибка скачивания {file_path}: {e}")
            return None
    
    def analyze_example_structure(self, file_path: str) -> Dict[str, Any]:
        """Анализирует структуру примера документа"""
        try:
            from document_analyzer import DocumentAnalyzer
            analyzer = DocumentAnalyzer()
            return analyzer.analyze_document_structure(file_path)
        except Exception as e:
            print(f"❌ Ошибка анализа {file_path}: {e}")
            return {}
    
    def generate_based_on_supabase_examples(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Генерирует документ на основе примеров из Supabase Storage"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        # Получаем примеры из Supabase
        examples = self.get_examples_from_supabase(template_type)
        
        if not examples:
            print(f"❌ Не найдено примеров для типа {template_type} в Supabase Storage")
            return self._generate_fallback_document(template_type, command_data)
        
        # Скачиваем первый пример для анализа
        example_path = examples[0]
        temp_file = self.download_example_from_supabase(example_path)
        
        if not temp_file:
            print(f"❌ Не удалось скачать пример {example_path}")
            return self._generate_fallback_document(template_type, command_data)
        
        # Анализируем структуру примера
        example_structure = self.analyze_example_structure(temp_file)
        
        print(f"📚 Используем пример из Supabase: {os.path.basename(example_path)}")
        print(f"📊 Структура: {len(example_structure.get('structure', {}).get('headings', []))} заголовков, {len(example_structure.get('tables', []))} таблиц")
        
        # Создаем новый документ на основе примера
        doc = Document()
        
        # Применяем структуру из примера
        self._apply_example_structure(doc, example_structure, template_type, apartment_id)
        
        # Удаляем временный файл
        try:
            os.remove(temp_file)
        except:
            pass
        
        # Сохраняем новый документ
        filename = f"supabase_learning_{template_type}_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        print(f"✅ Создан документ на основе Supabase примера: {filename}")
        return filepath
    
    def _apply_example_structure(self, doc: Document, structure: Dict[str, Any], template_type: str, apartment_id: str):
        """Применяет структуру из примера к новому документу"""
        
        # Заголовок документа
        if template_type == 'handover_act':
            title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        elif template_type == 'defect_report':
            title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        elif template_type == 'work_report':
            title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        else:
            title = doc.add_heading('ДОКУМЕНТ', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Применяем структуру из примера
        headings = structure.get('structure', {}).get('headings', [])
        tables = structure.get('tables', [])
        
        # Добавляем секции на основе примера
        for heading in headings[:3]:  # Берем первые 3 заголовка
            if heading.get('text') and len(heading['text']) > 3:
                doc.add_heading(heading['text'], level=2)
                
                # Добавляем содержимое секции
                if 'информация' in heading['text'].lower():
                    self._add_apartment_info(doc, apartment_id)
                elif 'дефект' in heading['text'].lower():
                    self._add_defects_section(doc, apartment_id)
                elif 'работа' in heading['text'].lower():
                    self._add_works_section(doc, apartment_id)
                elif 'рекомендация' in heading['text'].lower():
                    self._add_recommendations_section(doc)
                else:
                    doc.add_paragraph(f'Содержимое секции "{heading["text"]}" будет добавлено здесь.')
        
        # Добавляем таблицы на основе примера
        for i, table_info in enumerate(tables[:2]):  # Берем первые 2 таблицы
            if table_info.get('headers'):
                doc.add_heading(f'Таблица {i+1}', level=2)
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=len(table_info['headers']))
                table.style = 'Table Grid'
                
                # Заголовки
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(table_info['headers']):
                    hdr_cells[j].text = header
                    # Жирный шрифт для заголовков
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Добавляем примеры данных
                if 'квартир' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = apartment_id
                    row_cells[1].text = 'Готов к приёмке' if i == 0 else '100%'
                elif 'дефект' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Проверка качества'
                    row_cells[1].text = 'Активен'
                elif 'работа' in ' '.join(table_info['headers']).lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Покраска стен'
                    row_cells[1].text = 'Завершено'
        
        # Подписи (стандартная секция)
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Заказчик: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        doc.add_paragraph('Технадзор: _________________')
    
    def _add_apartment_info(self, doc: Document, apartment_id: str):
        """Добавляет информацию о квартире"""
        doc.add_paragraph(f'Номер квартиры: {apartment_id}')
        doc.add_paragraph(f'Статус: Готов к приёмке')
        doc.add_paragraph(f'Площадь: 45.2 м²')
        doc.add_paragraph(f'Этаж: 11')
    
    def _add_defects_section(self, doc: Document, apartment_id: str):
        """Добавляет секцию дефектов"""
        doc.add_paragraph('Обнаруженные дефекты:')
        doc.add_paragraph('• Трещина в стене (координаты: X:150, Y:200)')
        doc.add_paragraph('• Неровность пола (координаты: X:300, Y:100)')
        doc.add_paragraph('• Проблемы с проводкой (координаты: X:50, Y:250)')
    
    def _add_works_section(self, doc: Document, apartment_id: str):
        """Добавляет секцию работ"""
        doc.add_paragraph('Выполненные работы:')
        doc.add_paragraph('• Покраска стен - завершено')
        doc.add_paragraph('• Укладка плитки - завершено')
        doc.add_paragraph('• Установка сантехники - в процессе')
    
    def _add_recommendations_section(self, doc: Document):
        """Добавляет секцию рекомендаций"""
        doc.add_paragraph('Рекомендации:')
        doc.add_paragraph('• Устранить обнаруженные дефекты')
        doc.add_paragraph('• Провести финальную проверку качества')
        doc.add_paragraph('• Подготовить документы для приёмки')
    
    def _generate_fallback_document(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Создает базовый документ если нет примеров"""
        apartment_id = command_data.get('apartment_id', 'Не указано')
        
        doc = Document()
        
        # Базовый заголовок
        if template_type == 'handover_act':
            title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        elif template_type == 'defect_report':
            title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        elif template_type == 'work_report':
            title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        else:
            title = doc.add_heading('ДОКУМЕНТ', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Базовая информация
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        doc.add_paragraph(f'Дата составления: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('Документ создан на основе базового шаблона.')
        
        # Сохраняем
        filename = f"fallback_{template_type}_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_professional_document(self, template_type: str, command_data: Dict[str, Any]) -> str:
        """Генерирует профессиональный документ на основе шаблона 7.docx"""
        apartment_id = command_data.get('apartment_id', '1101')
        
        # Скачиваем профессиональный шаблон
        temp_file = self.download_example_from_supabase('7.docx')
        
        if not temp_file:
            print("❌ Не удалось скачать профессиональный шаблон")
            return self._generate_fallback_document(template_type, command_data)
        
        # Открываем шаблон как документ
        doc = Document(temp_file)
        
        # Обновляем данные в документе
        self._update_professional_template(doc, template_type, apartment_id, command_data)
        
        # Удаляем временный файл
        try:
            os.remove(temp_file)
        except:
            pass
        
        # Сохраняем новый документ
        filename = f"professional_{template_type}_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        print(f"✅ Создан профессиональный документ: {filename}")
        return filepath
    
    def _update_professional_template(self, doc: Document, template_type: str, apartment_id: str, command_data: Dict[str, Any]):
        """Обновляет профессиональный шаблон с новыми данными"""
        from datetime import datetime
        
        # Обновляем дату
        current_date = datetime.now().strftime("%d.%m.%Yг.")
        current_time = datetime.now().strftime("%H:%M")
        
        # Получаем дополнительные данные из команды
        document_purpose = command_data.get('document_purpose', '')
        intent = command_data.get('meta', {}).get('intent', '')
        notes = command_data.get('meta', {}).get('notes', '')
        
        # Генерируем умное содержимое на основе запроса
        smart_content = self._generate_smart_content(template_type, apartment_id, command_data)
        print(f"🧠 Сгенерировано умное содержимое: {smart_content['main_content'][:100]}...")
        
        # Проходим по всем параграфам и обновляем данные
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Сначала умная адаптация содержимого под конкретный запрос
            if template_type == 'official_letter' or intent == 'official_letter':
                # Для писем заменяем весь основной параграф
                if "Прошу назначить ответственного за" in text or "приемку выполненных работ" in text:
                    print(f"🔄 Заменяем параграф: {text[:50]}...")
                    paragraph.text = smart_content['main_content']
                    print(f"✅ Новый параграф: {paragraph.text[:50]}...")
                    continue  # Пропускаем остальные обновления для этого параграфа
                
                # Адаптируем приветствие под получателя
                if "Уважаемый Иса Исаевич!" in text:
                    paragraph.text = smart_content['greeting']
                    continue
            else:
                # Для других типов документов
                if "приемку выполненных работ" in text:
                    if template_type == 'defect_report' or intent == 'defect_report':
                        defect_desc = command_data.get('defect_description', 'дефект обнаружен')
                        paragraph.text = text.replace("приемку выполненных работ", f"дефект: {defect_desc}")
                    elif template_type == 'handover_act' or intent == 'handover_act':
                        paragraph.text = text.replace("приемку выполненных работ", "приемку квартиры")
                    elif template_type == 'work_report' or intent == 'work_report':
                        paragraph.text = text.replace("приемку выполненных работ", "выполненные работы")
            
            # Обновляем дату в шаблоне 7.docx (только если параграф не был заменен)
            if "03.09.2025г." in text:
                paragraph.text = text.replace("03.09.2025г.", current_date)
            
            # Обновляем номер квартиры в списке квартир (только если параграф не был заменен)
            if "квартирах 401,  404, 501, 601, 603, 604, 804, 1004, 1204" in text:
                paragraph.text = text.replace("квартирах 401,  404, 501, 601, 603, 604, 804, 1004, 1204", f"квартире {apartment_id}")
            
            # Обновляем время (только если параграф не был заменен)
            if "в 11:00" in text:
                paragraph.text = text.replace("в 11:00", f"в {current_time}")
    
    def _generate_smart_content(self, template_type: str, apartment_id: str, command_data: Dict[str, Any]) -> Dict[str, str]:
        """Генерирует умное содержимое на основе запроса"""
        recipient = command_data.get('recipient', 'заказчику')
        notes = command_data.get('meta', {}).get('notes', '')
        defect_desc = command_data.get('defect_description', '')
        
        # Определяем тип получателя
        if 'заказчик' in recipient:
            greeting = "Уважаемый Заказчик!"
            company_name = "ООО «АВ Development»"
        elif 'подрядчик' in recipient:
            greeting = "Уважаемый Подрядчик!"
            company_name = "ООО «Интербилдинг»"
        elif 'руководитель' in recipient:
            greeting = "Уважаемый Руководитель!"
            company_name = "ООО «АВ Development»"
        else:
            greeting = "Уважаемый Заказчик!"
            company_name = "ООО «АВ Development»"
        
        # Генерируем содержимое на основе типа документа и запроса
        if template_type == 'official_letter' or command_data.get('meta', {}).get('intent') == 'official_letter':
            if 'дефект' in notes.lower() or defect_desc:
                main_content = f"Уведомляем Вас о выявленном дефекте в квартире {apartment_id} объекта по адресу: г. Москва, ЗАО, район Раменки, ул. Мосфильмовская, з/у 1В, корп. Т. {defect_desc if defect_desc else 'Дефект обнаружен'} требует немедленного устранения. Просим назначить ответственного за устранение дефекта и уведомить о сроках выполнения работ."
            elif 'акт' in notes.lower() or 'приёмк' in notes.lower():
                main_content = f"Уведомляем Вас о необходимости проведения приёмки квартиры {apartment_id} объекта по адресу: г. Москва, ЗАО, район Раменки, ул. Мосфильмовская, з/у 1В, корп. Т. Просим назначить ответственного за приёмку и согласовать дату и время проведения работ."
            elif 'отчёт' in notes.lower() or 'отчет' in notes.lower():
                main_content = f"Представляем отчёт о выполненных работах в квартире {apartment_id} объекта по адресу: г. Москва, ЗАО, район Раменки, ул. Мосфильмовская, з/у 1В, корп. Т. Все работы выполнены в соответствии с техническими требованиями и готовы к приёмке."
            else:
                main_content = f"Уведомляем Вас о необходимости проведения работ в квартире {apartment_id} объекта по адресу: г. Москва, ЗАО, район Раменки, ул. Мосфильмовская, з/у 1В, корп. Т. Просим назначить ответственного за координацию работ и уведомить о дальнейших действиях."
        else:
            # Для других типов документов
            main_content = f"Прошу назначить ответственного за {template_type} в квартире {apartment_id} объекта по адресу: г. Москва, ЗАО, район Раменки, ул. Мосфильмовская, з/у 1В, корп. Т."
        
        return {
            'greeting': greeting,
            'main_content': main_content,
            'company_name': company_name
        }

def main():
    """Тестирование генератора с Supabase"""
    generator = SupabaseLearningGenerator()
    
    # Тестовые данные
    test_data = {
        'apartment_id': '1101',
        'act_type': 'handover'
    }
    
    print("🧪 Тестируем генератор с Supabase Storage")
    
    # Тестируем разные типы
    for doc_type in ['handover_act', 'defect_report', 'work_report']:
        print(f"\n📄 Создаем {doc_type} на основе Supabase примеров...")
        result = generator.generate_based_on_supabase_examples(doc_type, test_data)
        print(f"✅ Результат: {result}")

if __name__ == "__main__":
    main()
