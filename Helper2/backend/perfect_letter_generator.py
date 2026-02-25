#!/usr/bin/env python3
"""
Идеальный генератор писем, точно копирующий структуру первого письма
"""

import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class PerfectLetterGenerator:
    def __init__(self, documents_dir: str = "../existing_documents"):
        self.documents_dir = documents_dir
        self.logo_path = "logo_image1.png"
        os.makedirs(documents_dir, exist_ok=True)
    
    def create_perfect_letter(self, apartment_id: str, issue_type: str, issue_description: str) -> str:
        """Создает письмо, точно копируя структуру первого письма"""
        
        # Создаем документ
        doc = Document()
        
        # Устанавливаем точные поля как в первом письме
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)      # 0.50 дюйма
            section.bottom_margin = Inches(0.5)   # 0.50 дюйма
            section.left_margin = Inches(0.5)     # 0.50 дюйма
            section.right_margin = Inches(0.5)    # 0.50 дюйма
        
        # 1. ТАБЛИЦА С ЛОГОТИПОМ И АДРЕСАМИ (точно как в первом письме)
        self._add_perfect_logo_table(doc)
        
        # 2. ОБРАЩЕНИЕ (по центру)
        self._add_perfect_greeting(doc)
        
        # 3. ОСНОВНОЙ ТЕКСТ (с точными отступами как в первом письме)
        self._add_perfect_main_content(doc, apartment_id, issue_type, issue_description)
        
        # 4. ПРОСЬБА О СОДЕЙСТВИИ (с отступами)
        self._add_perfect_request(doc)
        
        # 5. ПОДПИСЬ (точно как в первом письме)
        self._add_perfect_signature(doc)
        
        # Сохраняем документ
        filename = f"Исх. письмо {issue_type} {apartment_id} от {datetime.now().strftime('%d.%m.%y')}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_perfect_logo_table(self, doc: Document):
        """Добавляет таблицу точно как в первом письме"""
        # Создаем таблицу 1x2
        table = doc.add_table(rows=1, cols=2)
        table.style = None  # Без рамок
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Получаем ячейки
        cells = table.rows[0].cells
        
        # Левая ячейка - логотип и адрес отправителя (по центру как в оригинале)
        left_cell = cells[0]
        
        # Добавляем логотип, если он существует
        if os.path.exists(self.logo_path):
            try:
                paragraph = left_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # По центру как в оригинале
                
                # Добавляем адрес под логотипом
                address_para = left_cell.add_paragraph()
                address_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
                address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # По центру как в оригинале
                
            except Exception as e:
                print(f"Ошибка добавления логотипа: {e}")
                left_para = left_cell.paragraphs[0]
                left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
                left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            left_para = left_cell.paragraphs[0]
            left_para.text = "Российская Федерация, 124498, Россия, Москва г.,\nЗеленоград г., 4922-й проезд, строение 2"
            left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Правая ячейка - адрес получателя (по правому краю как в оригинале)
        right_cell = cells[1]
        
        # Создаем отдельные абзацы для каждой строки адресата
        recipient1 = right_cell.add_paragraph()
        recipient1.text = "Руководителю проекта"
        recipient1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        recipient2 = right_cell.add_paragraph()
        recipient2.text = "ООО «АВ Development»"
        recipient2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        recipient3 = right_cell.add_paragraph()
        recipient3.text = "Эльман И.И."
        recipient3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Добавляем пустую строку после таблицы
        doc.add_paragraph()
    
    def _add_perfect_greeting(self, doc: Document):
        """Добавляет обращение точно как в первом письме"""
        greeting = doc.add_paragraph()
        greeting.text = "Уважаемый Иса Исаевич!"
        greeting.alignment = WD_ALIGN_PARAGRAPH.CENTER  # По центру как в оригинале
        
        # Добавляем пустую строку
        doc.add_paragraph()
    
    def _add_perfect_main_content(self, doc: Document, apartment_id: str, issue_type: str, issue_description: str):
        """Добавляет основной текст с точными отступами как в первом письме"""
        
        # Первый абзац - с левым отступом 14.2 pt и отступом первой строки 28.35 pt
        start_para = doc.add_paragraph()
        start_para.text = f"Уведомляем Вас о том, что {issue_description.lower()}"
        start_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        start_para.paragraph_format.left_indent = Pt(14.2)
        start_para.paragraph_format.first_line_indent = Pt(28.35)
        
        # Дополнительная информация в зависимости от типа проблемы
        if "смещение сроков" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_para.text = "Данный факт влияет на сроки производства работ и монтаж инженерных систем, в том числе системы отопления, водоснабжения, вентиляции и электроснабжения компанией ООО «Интербилдинг» - сроки будут увеличены."
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            additional_para.paragraph_format.first_line_indent = Pt(35.45)  # Как в оригинале
        elif "дефект" in issue_type.lower() or "проблема" in issue_type.lower():
            additional_para = doc.add_paragraph()
            additional_para.text = f"Обнаруженные дефекты в квартире {apartment_id} требуют немедленного устранения для обеспечения качества выполняемых работ и соблюдения сроков сдачи объекта."
            additional_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            additional_para.paragraph_format.first_line_indent = Pt(35.45)  # Как в оригинале
        
        # Добавляем пустую строку
        doc.add_paragraph()
    
    def _add_perfect_request(self, doc: Document):
        """Добавляет просьбу о содействии с точными отступами как в первом письме"""
        request_para = doc.add_paragraph()
        request_para.text = "Просим Вас посодействовать в решении данного вопроса для ускорения процесса сдачи и передачи инженерных систем компанией ООО «Сварго» и выполнения монтажных работ инженерных систем компанией ООО «Интербилдинг»."
        request_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        request_para.paragraph_format.first_line_indent = Pt(35.45)  # Как в оригинале
        request_para.paragraph_format.line_spacing = 1.0  # Межстрочный интервал как в оригинале
        
        # Добавляем несколько пустых строк перед подписью
        for _ in range(8):
            doc.add_paragraph()
    
    def _add_perfect_signature(self, doc: Document):
        """Добавляет подпись точно как в первом письме"""
        # Должность - с отступом слева (как в оригинале)
        position_para1 = doc.add_paragraph()
        position_para1.text = "           Заместитель директора и"
        position_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        position_para1.paragraph_format.line_spacing = 1.0  # Как в оригинале
        
        position_para2 = doc.add_paragraph()
        position_para2.text = "           руководитель проекта строительства"
        position_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        position_para2.paragraph_format.line_spacing = 1.0  # Как в оригинале
        
        # Компания слева, подпись ближе к правому краю
        signature_para = doc.add_paragraph()
        signature_para.text = "           ООО «Интербилдинг»"
        signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_para.paragraph_format.line_spacing = 1.0  # Как в оригинале
        
        # Отдельный абзац для подписи по правому краю
        signature_name_para = doc.add_paragraph()
        signature_name_para.text = "Кучун Р.В."
        signature_name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_name_para.paragraph_format.line_spacing = 1.0

def create_perfect_letter():
    """Создает идеальное письмо, точно копируя первое письмо"""
    print("📧 Создание идеального письма по образцу...")
    
    generator = PerfectLetterGenerator()
    
    # Тестовые данные
    apartment_id = "2301"
    issue_type = "смещение сроков поставки оборудования"
    issue_description = "задержка в поставке сантехнического оборудования для квартиры 2301, что влияет на общие сроки завершения работ"
    
    try:
        result = generator.create_perfect_letter(
            apartment_id=apartment_id,
            issue_type=issue_type,
            issue_description=issue_description
        )
        
        if result:
            print("✅ Идеальное письмо создано!")
            print(f"📁 Файл: {result}")
            
            # Показываем содержимое
            print("\n📄 СОДЕРЖИМОЕ ИДЕАЛЬНОГО ПИСЬМА:")
            print("=" * 60)
            
            from docx import Document
            doc = Document(result)
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    print(f"{i:2d}. {paragraph.text}")
            
            print("=" * 60)
            
            return result
        else:
            print("❌ Не удалось создать идеальное письмо")
            return None
            
    except Exception as e:
        print(f"❌ Ошибка при создании идеального письма: {e}")
        return None

if __name__ == "__main__":
    create_perfect_letter()
