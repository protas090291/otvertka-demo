import os
import uuid
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json

class DocumentGenerator:
    def __init__(self, documents_dir: str = "documents"):
        self.documents_dir = documents_dir
        os.makedirs(documents_dir, exist_ok=True)
    
    def generate_handover_act(self, command_data: Dict[str, Any]) -> str:
        """Генерирует акт приёмки квартиры"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('АКТ ПРИЁМКИ КВАРТИРЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Основная информация
        doc.add_heading('Информация о квартире', level=2)
        
        # Таблица с данными
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Заполняем таблицу
        data = [
            ('Номер квартиры:', apartment_id),
            ('Тип акта:', command_data.get('act_type', 'Приёмка')),
            ('Статус:', 'Готов к приёмке'),
            ('Ответственный:', 'Технадзор'),
            ('Дата проверки:', now.strftime("%d.%m.%Y")),
            ('Примечания:', command_data.get('meta', {}).get('notes', 'Без замечаний'))
        ]
        
        for i, (key, value) in enumerate(data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            # Жирный шрифт для ключей
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Заказчик: _________________')
        doc.add_paragraph('Подрядчик: _________________')
        doc.add_paragraph('Технадзор: _________________')
        
        # Сохраняем документ
        filename = f"act_handover_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_defect_report(self, command_data: Dict[str, Any]) -> str:
        """Генерирует отчет о дефектах"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ДЕФЕКТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Описание дефекта
        doc.add_heading('Описание дефекта', level=2)
        defect_description = command_data.get('defect_description', 'Дефект не описан')
        doc.add_paragraph(defect_description)
        
        # Дополнительная информация
        doc.add_heading('Дополнительная информация', level=2)
        
        # Таблица с данными
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        data = [
            ('Номер квартиры:', apartment_id),
            ('Тип дефекта:', 'Строительный'),
            ('Приоритет:', 'Средний'),
            ('Статус:', 'Обнаружен')
        ]
        
        for i, (key, value) in enumerate(data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            # Жирный шрифт для ключей
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        doc.add_paragraph('1. Устранить дефект в течение 7 дней')
        doc.add_paragraph('2. Провести повторную проверку')
        doc.add_paragraph('3. Уведомить заказчика о результатах')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Обнаружил: _________________')
        doc.add_paragraph('Ответственный: _________________')
        
        # Сохраняем документ
        filename = f"defect_report_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_work_report(self, command_data: Dict[str, Any]) -> str:
        """Генерирует отчет о выполненных работах"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ОТЧЕТ О ВЫПОЛНЕННЫХ РАБОТАХ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о квартире
        apartment_id = command_data.get('apartment_id', 'Не указано')
        doc.add_heading(f'Квартира № {apartment_id}', level=1)
        
        # Дата и время
        now = datetime.now()
        doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y %H:%M")}')
        
        # Выполненные работы
        doc.add_heading('Выполненные работы', level=2)
        
        # Таблица работ
        table = doc.add_table(rows=5, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        headers = ['Наименование работы', 'Статус', 'Дата выполнения']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            for paragraph in table.cell(0, i).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Данные работ
        works = [
            ('Подготовка стен', 'Выполнено', now.strftime("%d.%m.%Y")),
            ('Покраска', 'Выполнено', now.strftime("%d.%m.%Y")),
            ('Укладка полов', 'В процессе', 'В работе'),
            ('Установка сантехники', 'Запланировано', 'Планируется')
        ]
        
        for i, (work, status, date) in enumerate(works, 1):
            table.cell(i, 0).text = work
            table.cell(i, 1).text = status
            table.cell(i, 2).text = date
        
        # Прогресс
        doc.add_heading('Общий прогресс', level=2)
        doc.add_paragraph('Выполнено: 75%')
        doc.add_paragraph('Осталось: 25%')
        
        # Подписи
        doc.add_heading('Подписи', level=2)
        doc.add_paragraph('Прораб: _________________')
        doc.add_paragraph('Технадзор: _________________')
        
        # Сохраняем документ
        filename = f"work_report_{apartment_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(self.documents_dir, filename)
        doc.save(filepath)
        
        return filepath

