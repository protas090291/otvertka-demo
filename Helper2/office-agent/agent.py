"""
Офисный агент для обработки команд голосового помощника
Читает команды из Supabase, генерирует документы и отправляет на печать
"""

import os
import sys
import time
import json
import logging
import asyncio
from datetime import datetime, timezone
from typing import Dict, Any, Optional, List
import traceback

# Импорты для работы с Supabase
from supabase import create_client, Client

# Импорты для генерации документов
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

# Импорты для работы с изображениями
from PIL import Image
import requests

# Импорты для печати (Windows)
try:
    import win32api
    import win32print
    WINDOWS_PRINTING = True
except ImportError:
    WINDOWS_PRINTING = False
    print("Windows printing libraries not available. Install pywin32 for Windows printing support.")

# Импорты для печати (Linux/macOS)
import subprocess
import platform

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('agent.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class OfficeAgent:
    """Офисный агент для обработки команд"""
    
    def __init__(self):
        self.supabase: Optional[Client] = None
        self.running = False
        self.poll_interval = 5  # секунд
        self.max_retries = 3
        self.printer_name = os.getenv("PRINTER_NAME", "")
        
        # Инициализация Supabase
        self._init_supabase()
        
        # Проверка принтера
        self._check_printer()
    
    def _init_supabase(self):
        """Инициализация Supabase клиента"""
        try:
            url = os.getenv("SUPABASE_URL")
            service_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
            
            if not url or not service_key:
                raise ValueError("SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY must be set")
            
            self.supabase = create_client(url, service_key)
            logger.info("Supabase client initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize Supabase client: {e}")
            raise
    
    def _check_printer(self):
        """Проверка доступности принтера"""
        try:
            if WINDOWS_PRINTING and platform.system() == "Windows":
                printers = [printer[2] for printer in win32print.EnumPrinters(2)]
                if self.printer_name and self.printer_name not in printers:
                    logger.warning(f"Printer '{self.printer_name}' not found. Available printers: {printers}")
                elif not self.printer_name:
                    default_printer = win32print.GetDefaultPrinter()
                    self.printer_name = default_printer
                    logger.info(f"Using default printer: {default_printer}")
            else:
                # Linux/macOS
                try:
                    result = subprocess.run(['lpstat', '-p'], capture_output=True, text=True)
                    if result.returncode == 0:
                        logger.info("CUPS printing system available")
                    else:
                        logger.warning("CUPS printing system not available")
                except FileNotFoundError:
                    logger.warning("lpstat command not found. Printing may not work on this system.")
                    
        except Exception as e:
            logger.error(f"Error checking printer: {e}")
    
    def _log_agent_action(self, command_id: str, level: str, message: str, details: Dict = None):
        """Логирование действий агента"""
        try:
            log_data = {
                "command_id": command_id,
                "level": level,
                "message": message,
                "details": details or {}
            }
            
            result = self.supabase.table("agent_logs").insert(log_data).execute()
            logger.info(f"Logged action: {level} - {message}")
            
        except Exception as e:
            logger.error(f"Failed to log agent action: {e}")
    
    def _update_command_status(self, command_id: str, status: str, result_url: str = None, error_message: str = None):
        """Обновление статуса команды"""
        try:
            update_data = {
                "status": status,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
            
            if status in ["done", "failed"]:
                update_data["processed_at"] = datetime.now(timezone.utc).isoformat()
            
            if result_url:
                update_data["result_url"] = result_url
                
            if error_message:
                update_data["error_message"] = error_message
            
            result = self.supabase.table("commands").update(update_data).eq("id", command_id).execute()
            
            if result.data:
                logger.info(f"Command {command_id} status updated to {status}")
                return True
            else:
                logger.error(f"Failed to update command {command_id} status")
                return False
                
        except Exception as e:
            logger.error(f"Error updating command status: {e}")
            return False
    
    def _get_pending_commands(self, limit: int = 10) -> List[Dict]:
        """Получение pending команд"""
        try:
            result = self.supabase.table("commands").select("*").eq("status", "pending").order("created_at").limit(limit).execute()
            return result.data or []
            
        except Exception as e:
            logger.error(f"Error fetching pending commands: {e}")
            return []
    
    def _increment_attempt_count(self, command_id: str):
        """Увеличение счетчика попыток"""
        try:
            result = self.supabase.table("commands").select("attempt_count").eq("id", command_id).execute()
            if result.data:
                current_count = result.data[0].get("attempt_count", 0)
                self.supabase.table("commands").update({
                    "attempt_count": current_count + 1
                }).eq("id", command_id).execute()
                
        except Exception as e:
            logger.error(f"Error incrementing attempt count: {e}")
    
    def _create_handover_act(self, payload: Dict[str, Any]) -> str:
        """Создание акта приёмки"""
        try:
            apartment_id = payload.get("apartment_id")
            act_type = payload.get("act_type", "handover")
            notes = payload.get("meta", {}).get("notes", "")
            
            # Создаем документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('Акт приёмки квартиры', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация о квартире
            doc.add_heading('Информация о квартире', level=1)
            
            # Таблица с информацией
            table = doc.add_table(rows=4, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Заполняем таблицу
            table.cell(0, 0).text = 'Номер квартиры:'
            table.cell(0, 1).text = str(apartment_id)
            
            table.cell(1, 0).text = 'Дата составления:'
            table.cell(1, 1).text = datetime.now().strftime('%d.%m.%Y')
            
            table.cell(2, 0).text = 'Тип акта:'
            table.cell(2, 1).text = 'Приёмка' if act_type == 'handover' else 'Дефектный'
            
            table.cell(3, 0).text = 'Примечания:'
            table.cell(3, 1).text = notes or 'Нет'
            
            # Подписи
            doc.add_heading('Подписи', level=1)
            doc.add_paragraph('Представитель застройщика: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Представитель заказчика: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Дата: _________________')
            
            # Сохраняем документ
            filename = f"handover_act_{apartment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("generated_docs", filename)
            
            # Создаем папку если не существует
            os.makedirs("generated_docs", exist_ok=True)
            
            doc.save(filepath)
            logger.info(f"Handover act created: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating handover act: {e}")
            raise
    
    def _create_defect_report(self, payload: Dict[str, Any]) -> str:
        """Создание отчёта о дефектах"""
        try:
            apartment_id = payload.get("apartment_id")
            defect_description = payload.get("defect_description", "")
            notes = payload.get("meta", {}).get("notes", "")
            
            # Создаем документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('Отчёт о дефектах', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация
            doc.add_heading('Информация о дефекте', level=1)
            
            table = doc.add_table(rows=3, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = 'Номер квартиры:'
            table.cell(0, 1).text = str(apartment_id)
            
            table.cell(1, 0).text = 'Дата обнаружения:'
            table.cell(1, 1).text = datetime.now().strftime('%d.%m.%Y')
            
            table.cell(2, 0).text = 'Описание дефекта:'
            table.cell(2, 1).text = defect_description or 'Не указано'
            
            # Дополнительные примечания
            if notes:
                doc.add_heading('Дополнительные примечания', level=1)
                doc.add_paragraph(notes)
            
            # Подписи
            doc.add_heading('Подписи', level=1)
            doc.add_paragraph('Обнаружил дефект: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Принял к сведению: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Дата: _________________')
            
            # Сохраняем документ
            filename = f"defect_report_{apartment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("generated_docs", filename)
            
            os.makedirs("generated_docs", exist_ok=True)
            doc.save(filepath)
            logger.info(f"Defect report created: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating defect report: {e}")
            raise
    
    def _upload_document_to_storage(self, filepath: str, apartment_id: str, doc_type: str) -> str:
        """Загрузка документа в Supabase Storage"""
        try:
            filename = os.path.basename(filepath)
            storage_path = f"{apartment_id}/{doc_type}/{filename}"
            
            with open(filepath, 'rb') as f:
                result = self.supabase.storage.from_("documents").upload(storage_path, f)
            
            if result:
                # Получаем публичный URL
                public_url = self.supabase.storage.from_("documents").get_public_url(storage_path)
                
                # Сохраняем информацию о документе в базу
                doc_data = {
                    "apartment_id": apartment_id,
                    "doc_type": doc_type,
                    "storage_path": storage_path,
                    "file_name": filename,
                    "file_size": os.path.getsize(filepath),
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
                
                self.supabase.table("documents").insert(doc_data).execute()
                
                logger.info(f"Document uploaded to storage: {public_url}")
                return public_url
            else:
                raise Exception("Failed to upload document to storage")
                
        except Exception as e:
            logger.error(f"Error uploading document to storage: {e}")
            raise
    
    def _print_document(self, filepath: str) -> bool:
        """Печать документа"""
        try:
            if WINDOWS_PRINTING and platform.system() == "Windows":
                # Windows печать
                win32api.ShellExecute(
                    0,
                    "print",
                    filepath,
                    f'/d:"{self.printer_name}"',
                    ".",
                    0
                )
                logger.info(f"Document sent to Windows printer: {self.printer_name}")
                return True
                
            else:
                # Linux/macOS печать через CUPS
                cmd = ['lp', '-d', self.printer_name, filepath] if self.printer_name else ['lp', filepath]
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                if result.returncode == 0:
                    logger.info(f"Document sent to printer: {self.printer_name or 'default'}")
                    return True
                else:
                    logger.error(f"Print command failed: {result.stderr}")
                    return False
                    
        except Exception as e:
            logger.error(f"Error printing document: {e}")
            return False
    
    def _process_command(self, command: Dict[str, Any]) -> bool:
        """Обработка одной команды"""
        command_id = command['id']
        command_type = command['type']
        payload = command['payload']
        
        try:
            logger.info(f"Processing command {command_id} of type {command_type}")
            self._log_agent_action(command_id, "info", f"Started processing command {command_type}")
            
            # Обновляем статус на "processing"
            self._update_command_status(command_id, "processing")
            
            result_url = None
            
            if command_type == "create_act":
                # Создание акта приёмки
                filepath = self._create_handover_act(payload)
                result_url = self._upload_document_to_storage(filepath, payload['apartment_id'], "handover_act")
                
            elif command_type == "print_act":
                # Печать акта
                # Сначала создаем акт, если его нет
                filepath = self._create_handover_act(payload)
                result_url = self._upload_document_to_storage(filepath, payload['apartment_id'], "handover_act")
                
                # Печатаем документ
                print_success = self._print_document(filepath)
                if not print_success:
                    raise Exception("Failed to print document")
                    
            elif command_type == "create_defect":
                # Создание отчёта о дефектах
                filepath = self._create_defect_report(payload)
                result_url = self._upload_document_to_storage(filepath, payload['apartment_id'], "defect_report")
                
            elif command_type == "print_defect_report":
                # Печать отчёта о дефектах
                filepath = self._create_defect_report(payload)
                result_url = self._upload_document_to_storage(filepath, payload['apartment_id'], "defect_report")
                
                # Печатаем документ
                print_success = self._print_document(filepath)
                if not print_success:
                    raise Exception("Failed to print document")
            
            # Обновляем статус на "done"
            self._update_command_status(command_id, "done", result_url)
            self._log_agent_action(command_id, "info", f"Command {command_type} completed successfully", {
                "result_url": result_url
            })
            
            logger.info(f"Command {command_id} processed successfully")
            return True
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error processing command {command_id}: {error_msg}")
            logger.error(traceback.format_exc())
            
            # Обновляем статус на "failed"
            self._update_command_status(command_id, "failed", error_message=error_msg)
            self._log_agent_action(command_id, "error", f"Command {command_type} failed", {
                "error": error_msg,
                "traceback": traceback.format_exc()
            })
            
            return False
    
    def _should_retry_command(self, command: Dict[str, Any]) -> bool:
        """Проверка, нужно ли повторить команду"""
        attempt_count = command.get('attempt_count', 0)
        return attempt_count < self.max_retries
    
    def run_once(self):
        """Однократный запуск обработки команд"""
        try:
            pending_commands = self._get_pending_commands()
            
            if not pending_commands:
                logger.debug("No pending commands found")
                return
            
            logger.info(f"Found {len(pending_commands)} pending commands")
            
            for command in pending_commands:
                command_id = command['id']
                attempt_count = command.get('attempt_count', 0)
                
                # Проверяем, не превышено ли количество попыток
                if attempt_count >= self.max_retries:
                    logger.warning(f"Command {command_id} exceeded max retries ({self.max_retries})")
                    self._update_command_status(command_id, "failed", error_message="Max retries exceeded")
                    continue
                
                # Увеличиваем счетчик попыток
                self._increment_attempt_count(command_id)
                
                # Обрабатываем команду
                success = self._process_command(command)
                
                if success:
                    logger.info(f"Command {command_id} processed successfully")
                else:
                    logger.error(f"Command {command_id} processing failed")
                    
        except Exception as e:
            logger.error(f"Error in run_once: {e}")
            logger.error(traceback.format_exc())
    
    def start(self):
        """Запуск агента в режиме polling"""
        logger.info("Starting office agent...")
        self.running = True
        
        try:
            while self.running:
                self.run_once()
                time.sleep(self.poll_interval)
                
        except KeyboardInterrupt:
            logger.info("Agent stopped by user")
        except Exception as e:
            logger.error(f"Agent error: {e}")
            logger.error(traceback.format_exc())
        finally:
            self.running = False
            logger.info("Agent stopped")
    
    def stop(self):
        """Остановка агента"""
        logger.info("Stopping agent...")
        self.running = False

def main():
    """Главная функция"""
    # Проверяем переменные окружения
    required_env_vars = ["SUPABASE_URL", "SUPABASE_SERVICE_ROLE_KEY"]
    missing_vars = [var for var in required_env_vars if not os.getenv(var)]
    
    if missing_vars:
        logger.error(f"Missing required environment variables: {missing_vars}")
        sys.exit(1)
    
    # Создаем и запускаем агента
    agent = OfficeAgent()
    
    try:
        agent.start()
    except Exception as e:
        logger.error(f"Failed to start agent: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()


