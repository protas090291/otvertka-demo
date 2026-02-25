"""
Упрощенный офисный агент для демонстрации
Работает с demo backend без Supabase
"""

import os
import sys
import time
import json
import logging
import requests
from datetime import datetime, timezone
from typing import Dict, Any, Optional, List
import traceback

# Импорты для генерации документов
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('simple_agent.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

class SimpleOfficeAgent:
    """Упрощенный офисный агент для демонстрации"""
    
    def __init__(self):
        self.api_base_url = "http://localhost:8000"
        self.running = False
        self.poll_interval = 5  # секунд
        self.max_retries = 3
        
        # Создаем папку для документов
        os.makedirs("generated_docs", exist_ok=True)
        
        logger.info("Simple Office Agent initialized")
    
    def _log_agent_action(self, command_id: str, level: str, message: str, details: Dict = None):
        """Логирование действий агента"""
        logger.info(f"[{level.upper()}] {message} - Command: {command_id}")
        if details:
            logger.info(f"Details: {details}")
    
    def _get_pending_commands(self) -> List[Dict]:
        """Получение pending команд из API"""
        try:
            response = requests.get(f"{self.api_base_url}/api/commands/pending")
            
            if response.status_code == 200:
                commands = response.json()
                logger.info(f"Found {len(commands)} pending commands")
                return commands
            else:
                logger.error(f"Error fetching pending commands: {response.status_code}")
                return []
                
        except Exception as e:
            logger.error(f"Error fetching pending commands: {e}")
            return []
    
    def _update_command_status(self, command_id: str, status: str, result_url: str = None, error_message: str = None):
        """Обновление статуса команды через API"""
        try:
            update_data = {"status": status}
            
            if result_url:
                update_data["result_url"] = result_url
                
            if error_message:
                update_data["error_message"] = error_message
            
            response = requests.patch(
                f"{self.api_base_url}/api/commands/{command_id}",
                json=update_data
            )
            
            if response.status_code == 200:
                logger.info(f"Command {command_id} status updated to {status}")
                return True
            else:
                logger.error(f"Failed to update command {command_id} status: {response.status_code}")
                return False
                
        except Exception as e:
            logger.error(f"Error updating command status: {e}")
            return False
    
    def _create_handover_act(self, payload: Dict[str, Any]) -> str:
        """Создание акта приёмки"""
        try:
            apartment_id = payload.get("apartment_id")
            act_type = payload.get("act_type", "handover")
            notes = payload.get("meta", {}).get("notes", "")
            
            # Создаем документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('Акт приёмки квартиры', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация о квартире
            doc.add_heading('Информация о квартире', level=1)
            
            # Таблица с информацией
            table = doc.add_table(rows=4, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Заполняем таблицу
            table.cell(0, 0).text = 'Номер квартиры:'
            table.cell(0, 1).text = str(apartment_id)
            
            table.cell(1, 0).text = 'Дата составления:'
            table.cell(1, 1).text = datetime.now().strftime('%d.%m.%Y')
            
            table.cell(2, 0).text = 'Тип акта:'
            table.cell(2, 1).text = 'Приёмка' if act_type == 'handover' else 'Дефектный'
            
            table.cell(3, 0).text = 'Примечания:'
            table.cell(3, 1).text = notes or 'Нет'
            
            # Подписи
            doc.add_heading('Подписи', level=1)
            doc.add_paragraph('Представитель застройщика: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Представитель заказчика: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Дата: _________________')
            
            # Сохраняем документ
            filename = f"handover_act_{apartment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("generated_docs", filename)
            
            doc.save(filepath)
            logger.info(f"Handover act created: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating handover act: {e}")
            raise
    
    def _create_defect_report(self, payload: Dict[str, Any]) -> str:
        """Создание отчёта о дефектах"""
        try:
            apartment_id = payload.get("apartment_id")
            defect_description = payload.get("defect_description", "")
            notes = payload.get("meta", {}).get("notes", "")
            
            # Создаем документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('Отчёт о дефектах', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация
            doc.add_heading('Информация о дефекте', level=1)
            
            table = doc.add_table(rows=3, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = 'Номер квартиры:'
            table.cell(0, 1).text = str(apartment_id)
            
            table.cell(1, 0).text = 'Дата обнаружения:'
            table.cell(1, 1).text = datetime.now().strftime('%d.%m.%Y')
            
            table.cell(2, 0).text = 'Описание дефекта:'
            table.cell(2, 1).text = defect_description or 'Не указано'
            
            # Дополнительные примечания
            if notes:
                doc.add_heading('Дополнительные примечания', level=1)
                doc.add_paragraph(notes)
            
            # Подписи
            doc.add_heading('Подписи', level=1)
            doc.add_paragraph('Обнаружил дефект: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Принял к сведению: _________________')
            doc.add_paragraph('')
            doc.add_paragraph('Дата: _________________')
            
            # Сохраняем документ
            filename = f"defect_report_{apartment_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("generated_docs", filename)
            
            doc.save(filepath)
            logger.info(f"Defect report created: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error creating defect report: {e}")
            raise
    
    def _simulate_document_upload(self, filepath: str, apartment_id: str, doc_type: str) -> str:
        """Имитация загрузки документа в storage"""
        try:
            filename = os.path.basename(filepath)
            # Имитируем URL документа
            public_url = f"https://demo-storage.example.com/documents/{apartment_id}/{doc_type}/{filename}"
            
            logger.info(f"Document uploaded to storage: {public_url}")
            return public_url
                
        except Exception as e:
            logger.error(f"Error simulating document upload: {e}")
            raise
    
    def _simulate_printing(self, filepath: str) -> bool:
        """Имитация печати документа"""
        try:
            # В реальном приложении здесь была бы печать
            logger.info(f"Document sent to printer: {filepath}")
            logger.info("🖨️ [DEMO] Document would be printed here")
            return True
                    
        except Exception as e:
            logger.error(f"Error simulating printing: {e}")
            return False
    
    def _process_command(self, command: Dict[str, Any]) -> bool:
        """Обработка одной команды"""
        command_id = command['id']
        command_type = command['type']
        
        # Получаем полную информацию о команде
        try:
            response = requests.get(f"{self.api_base_url}/api/commands/{command_id}")
            if response.status_code == 200:
                full_command = response.json()
                payload = full_command['payload']
            else:
                logger.error(f"Failed to get full command details for {command_id}")
                return False
        except Exception as e:
            logger.error(f"Error getting full command details: {e}")
            return False
        
        try:
            logger.info(f"Processing command {command_id} of type {command_type}")
            self._log_agent_action(command_id, "info", f"Started processing command {command_type}")
            
            # Обновляем статус на "processing"
            self._update_command_status(command_id, "processing")
            
            result_url = None
            
            if command_type == "create_act":
                # Создание акта приёмки
                filepath = self._create_handover_act(payload)
                result_url = self._simulate_document_upload(filepath, payload['apartment_id'], "handover_act")
                
            elif command_type == "print_act":
                # Печать акта
                # Сначала создаем акт, если его нет
                filepath = self._create_handover_act(payload)
                result_url = self._simulate_document_upload(filepath, payload['apartment_id'], "handover_act")
                
                # Печатаем документ
                print_success = self._simulate_printing(filepath)
                if not print_success:
                    raise Exception("Failed to print document")
                    
            elif command_type == "create_defect":
                # Создание отчёта о дефектах
                filepath = self._create_defect_report(payload)
                result_url = self._simulate_document_upload(filepath, payload['apartment_id'], "defect_report")
                
            elif command_type == "print_defect_report":
                # Печать отчёта о дефектах
                filepath = self._create_defect_report(payload)
                result_url = self._simulate_document_upload(filepath, payload['apartment_id'], "defect_report")
                
                # Печатаем документ
                print_success = self._simulate_printing(filepath)
                if not print_success:
                    raise Exception("Failed to print document")
            
            # Обновляем статус на "done"
            self._update_command_status(command_id, "done", result_url)
            self._log_agent_action(command_id, "info", f"Command {command_type} completed successfully", {
                "result_url": result_url
            })
            
            logger.info(f"Command {command_id} processed successfully")
            return True
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error processing command {command_id}: {error_msg}")
            logger.error(traceback.format_exc())
            
            # Обновляем статус на "failed"
            self._update_command_status(command_id, "failed", error_message=error_msg)
            self._log_agent_action(command_id, "error", f"Command {command_type} failed", {
                "error": error_msg,
                "traceback": traceback.format_exc()
            })
            
            return False
    
    def run_once(self):
        """Однократный запуск обработки команд"""
        try:
            pending_commands = self._get_pending_commands()
            
            if not pending_commands:
                logger.debug("No pending commands found")
                return
            
            logger.info(f"Found {len(pending_commands)} pending commands")
            
            for command in pending_commands:
                command_id = command['id']
                attempt_count = command.get('attempt_count', 0)
                
                # Проверяем, не превышено ли количество попыток
                if attempt_count >= self.max_retries:
                    logger.warning(f"Command {command_id} exceeded max retries ({self.max_retries})")
                    self._update_command_status(command_id, "failed", error_message="Max retries exceeded")
                    continue
                
                # Обрабатываем команду
                success = self._process_command(command)
                
                if success:
                    logger.info(f"Command {command_id} processed successfully")
                else:
                    logger.error(f"Command {command_id} processing failed")
                    
        except Exception as e:
            logger.error(f"Error in run_once: {e}")
            logger.error(traceback.format_exc())
    
    def start(self):
        """Запуск агента в режиме polling"""
        logger.info("Starting simple office agent...")
        self.running = True
        
        try:
            while self.running:
                self.run_once()
                time.sleep(self.poll_interval)
                
        except KeyboardInterrupt:
            logger.info("Agent stopped by user")
        except Exception as e:
            logger.error(f"Agent error: {e}")
            logger.error(traceback.format_exc())
        finally:
            self.running = False
            logger.info("Agent stopped")
    
    def stop(self):
        """Остановка агента"""
        logger.info("Stopping agent...")
        self.running = False

def main():
    """Главная функция"""
    # Проверяем доступность API
    try:
        response = requests.get("http://localhost:8000/health")
        if response.status_code != 200:
            logger.error("Backend API is not available")
            sys.exit(1)
        logger.info("Backend API is available")
    except Exception as e:
        logger.error(f"Cannot connect to backend API: {e}")
        logger.error("Make sure the backend is running on http://localhost:8000")
        sys.exit(1)
    
    # Создаем и запускаем агента
    agent = SimpleOfficeAgent()
    
    try:
        agent.start()
    except Exception as e:
        logger.error(f"Failed to start agent: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
